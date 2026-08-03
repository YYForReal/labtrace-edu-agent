#!/usr/bin/env python3
"""
评分表注入脚本
将评分汇总表格注入到 DOCX 文档末尾
"""

import os
import sys
import json
import shutil
import argparse

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("请安装 python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def inject_score_table(doc, grading_result):
    """
    在文档末尾注入评分汇总表

    Args:
        doc: python-docx Document 对象
        grading_result: 评分结果字典
    """
    # 添加分页符
    doc.add_page_break()

    # 添加标题
    title = doc.add_paragraph()
    title_run = title.add_run("【AI 辅助批改报告】")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加学生信息
    info_para = doc.add_paragraph()
    student_name = grading_result.get("student_name", "未知")
    student_id = grading_result.get("student_id", "未知")
    total_score = grading_result.get("total_score", 0)

    run1 = info_para.add_run(f"学生姓名：{student_name}    ")
    run1.bold = True
    run2 = info_para.add_run(f"学号：{student_id}    ")
    run2.bold = True
    run3 = info_para.add_run(f"总分：{total_score}分")
    run3.bold = True
    run3.font.size = Pt(14)

    doc.add_paragraph()  # 空行

    # 创建表格
    criteria = grading_result.get("criterion_scores", [])
    table = doc.add_table(rows=1, cols=4)

    # 尝试设置表格样式
    try:
        table.style = "Light Grid Accent 1"
    except Exception:
        try:
            table.style = "Table Grid"
        except Exception:
            pass

    # 表头
    header_cells = table.rows[0].cells
    headers = ["评分项", "满分", "得分", "评注"]
    for i, header_text in enumerate(headers):
        header_cells[i].text = header_text
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 设置背景色
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "003366")
        header_cells[i]._tc.get_or_add_tcPr().append(shading)

    # 添加数据行
    for cs in criteria:
        row_cells = table.add_row().cells
        row_cells[0].text = cs.get("criterion_name", "")
        row_cells[1].text = str(cs.get("max_score", ""))
        row_cells[2].text = str(cs.get("score", ""))
        row_cells[3].text = cs.get("reason", "")

        # 居中对齐
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 颜色编码
        max_score = cs.get("max_score", 1)
        score = cs.get("score", 0)
        score_rate = score / max_score if max_score > 0 else 0

        if row_cells[2].paragraphs[0].runs:
            score_run = row_cells[2].paragraphs[0].runs[0]
            if score_rate >= 0.8:
                score_run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
            elif score_rate < 0.6:
                score_run.font.color.rgb = RGBColor(204, 0, 0)  # 红色

    # 总分行
    total_row = table.add_row().cells
    total_row[0].text = "总分"
    total_row[1].text = ""
    total_row[2].text = str(total_score)
    total_row[3].text = ""

    for i in range(4):
        total_row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if total_row[i].paragraphs[0].runs:
            total_row[i].paragraphs[0].runs[0].bold = True

    if total_row[2].paragraphs[0].runs:
        total_row[2].paragraphs[0].runs[0].font.size = Pt(14)

    doc.add_paragraph()


def main():
    parser = argparse.ArgumentParser(description="注入评分表到 DOCX")
    parser.add_argument("--input", "-i", required=True, help="输入 DOCX 文件")
    parser.add_argument("--grading-result", "-g", required=True, help="评分结果 JSON")
    parser.add_argument("--output", "-o", required=True, help="输出 DOCX 文件")

    args = parser.parse_args()

    with open(args.grading_result, "r", encoding="utf-8") as f:
        grading_result = json.load(f)

    # 复制原始文件
    if args.input != args.output:
        shutil.copy2(args.input, args.output)

    doc = Document(args.output)
    inject_score_table(doc, grading_result)
    doc.save(args.output)

    print(f"评分表已注入: {args.output}")


if __name__ == "__main__":
    main()
