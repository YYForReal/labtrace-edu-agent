#!/usr/bin/env python3
"""
助教总评注入脚本
将评语文本注入到文档末尾（新页）
"""

import sys
import argparse

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("请安装 python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def inject_final_comment(doc, comment_text):
    """
    在文档末尾新页注入助教总评

    Args:
        doc: python-docx Document 对象
        comment_text: 评语文本
    """
    # 新页
    doc.add_page_break()

    # 标题
    title = doc.add_paragraph()
    title_run = title.add_run("【助教总评】")
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 评语正文
    comment_para = doc.add_paragraph(comment_text)
    comment_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in comment_para.runs:
        run.font.size = Pt(12)

    # 设置段落间距
    fmt = comment_para.paragraph_format
    fmt.space_after = Pt(12)
    fmt.line_spacing = 1.5

    # 底部分隔线
    doc.add_paragraph()
    separator = doc.add_paragraph()
    sep_run = separator.add_run("─" * 40)
    sep_run.font.color.rgb = RGBColor(128, 128, 128)
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 标注
    note = doc.add_paragraph()
    note_run = note.add_run(
        "本报告由 AI 辅助批改系统生成，仅供参考，最终成绩以助教确认为准。"
    )
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(128, 128, 128)
    note_run.font.italic = True
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    parser = argparse.ArgumentParser(description="注入助教总评到 DOCX")
    parser.add_argument("--input", "-i", required=True, help="输入 DOCX 文件")
    parser.add_argument("--comment", "-c", required=True, help="评语文本或评语文件路径")
    parser.add_argument("--output", "-o", required=True, help="输出 DOCX 文件")

    args = parser.parse_args()

    # 读取评语
    import os

    if os.path.isfile(args.comment):
        with open(args.comment, "r", encoding="utf-8") as f:
            comment_text = f.read().strip()
    else:
        comment_text = args.comment

    doc = Document(args.input)
    inject_final_comment(doc, comment_text)
    doc.save(args.output)

    print(f"助教总评已注入: {args.output}")


if __name__ == "__main__":
    main()
