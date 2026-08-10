#!/usr/bin/env python3
"""
inject_grading_to_docx.py — 实验报告批改内容注入器（完整集成版）

功能概述：
  对学生提交的实验报告（.docx 格式）进行四类批改内容的定位注入：
  0. 引用批注 — 在文档正文指定位置添加 Word Comments
  1. 成绩评定 — 在表格最后一页的「成绩评定：」行填写分数公式
  2. 评语     — 在「评语：」行填写评语文本
  3. 签名+日期 — 在「指导教师签字：」后插入签名图片和日期

文档结构假设（基于深圳大学实验报告模板）：
  - 表格0: 封面信息（课程、学号、姓名等）
  - 表格1: 批改意见页，Cell[0,0] 包含：
    - Para[0]: "指导教师批阅意见："
    - Para[6]: "成绩评定： XX+XX+...=总分"
    - Para[8]: "评语：...."
    - Para[14]: "指导教师签字：" + 签名图片
    - Para[16]: "YYYY年 MM 月 DD 日"

用法：
  python inject_grading_to_docx.py \\
    --input student.docx \\
    --output graded.docx \\
    --config grading_config.json

  或在 Python 中导入使用：
    from inject_grading_to_docx import inject_all
"""

import argparse
import copy
import json
import os
import re
import sys
import zipfile
from datetime import datetime
from io import BytesIO
from lxml import etree

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Emu, RGBColor
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── 常量 ─────────────────────────────────────────────────
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
COMMENTS_CT = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)
COMMENTS_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)
BODY_HEADING_RE = re.compile(r"^\s*(?:[一二三四五六七八九十]+[、.．]|\d+[、.．]\s+)")
COVER_FOOTER_RE = re.compile(r"教\s*务\s*处\s*制")


# ══════════════════════════════════════════════════════════
# 模块0：引用批注注入
# ══════════════════════════════════════════════════════════


def _get_max_comment_id(doc):
    """获取文档中已有的最大 comment ID"""
    max_id = -1
    body = doc.element
    for elem in body.iter():
        tag = elem.tag
        if "commentRange" in tag or "commentReference" in tag:
            cid = elem.get(qn("w:id"))
            if cid is not None:
                max_id = max(max_id, int(cid))
    # 也检查 comments.xml
    try:
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        with zipfile.ZipFile(buf, "r") as zf:
            if "word/comments.xml" in zf.namelist():
                root = etree.fromstring(zf.read("word/comments.xml"))
                for c in root.findall(qn("w:comment")):
                    cid = c.get(qn("w:id"))
                    if cid is not None:
                        max_id = max(max_id, int(cid))
    except Exception:
        pass
    return max_id


def _paragraph_has_page_break(paragraph):
    """检测段落中是否包含显式分页符。"""
    para_elem = paragraph._element
    for br in para_elem.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return any(True for _ in para_elem.iter(qn("w:lastRenderedPageBreak")))


def _first_non_empty_paragraph_index(doc, start_index):
    """返回 start_index 之后首个非空段落索引。"""
    start_index = max(0, start_index)
    for i in range(start_index, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip():
            return i
    return len(doc.paragraphs)


def _find_annotation_start_index(doc):
    """
    计算批注允许定位的起点，默认跳过第一页封面。

    DOCX 没有稳定的“页”结构，因此优先使用显式分页符；若没有分页符，
    再按深圳大学实验报告模板的正文标题或“教务处制”尾标识别封面边界。
    """
    paragraphs = doc.paragraphs

    for i, paragraph in enumerate(paragraphs):
        if _paragraph_has_page_break(paragraph):
            return _first_non_empty_paragraph_index(doc, i + 1)

    for i, paragraph in enumerate(paragraphs):
        if BODY_HEADING_RE.match(paragraph.text.strip()):
            return i

    for i, paragraph in enumerate(paragraphs):
        if COVER_FOOTER_RE.search(paragraph.text):
            return _first_non_empty_paragraph_index(doc, i + 1)

    return 0


def _find_paragraph_by_keyword(doc, keyword, start_index=0):
    """按关键词在主体段落中查找，返回段落索引和元素"""
    normalized_keyword = re.sub(r"\s+", " ", keyword).strip()
    for i, p in enumerate(doc.paragraphs):
        if i < start_index:
            continue
        normalized_text = re.sub(r"\s+", " ", p.text).strip()
        if normalized_keyword in normalized_text:
            return i, p._element
    return -1, None


def _create_comment_xml(comment_id, author, date_str, initials, text):
    """创建单个 <w:comment> XML 元素"""
    comment = OxmlElement("w:comment")
    comment.set(qn("w:id"), str(comment_id))
    comment.set(qn("w:author"), author)
    comment.set(qn("w:date"), date_str)
    comment.set(qn("w:initials"), initials)

    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    p.append(pPr)

    # annotationRef run
    r_ref = OxmlElement("w:r")
    r_ref.append(OxmlElement("w:annotationRef"))
    p.append(r_ref)

    # 文本 run
    r_text = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:hint"), "eastAsia")
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "21")
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "24")
    rPr.append(szCs)

    lang = OxmlElement("w:lang")
    lang.set(qn("w:eastAsia"), "en-US")
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:bidi"), "ar-SA")
    rPr.append(lang)

    r_text.append(rPr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r_text.append(t)
    p.append(r_text)
    comment.append(p)
    return comment


def _mark_paragraph_with_comment(para_elem, comment_id_str):
    """在段落中插入 commentRangeStart/End/Reference 标记"""
    cid = comment_id_str

    range_start = OxmlElement("w:commentRangeStart")
    range_start.set(qn("w:id"), cid)

    range_end = OxmlElement("w:commentRangeEnd")
    range_end.set(qn("w:id"), cid)

    ref_run = OxmlElement("w:r")
    ref_rPr = OxmlElement("w:rPr")
    ref_style = OxmlElement("w:rStyle")
    ref_style.set(qn("w:val"), "CommentReference")
    ref_rPr.append(ref_style)
    ref_run.append(ref_rPr)
    comment_ref = OxmlElement("w:commentReference")
    comment_ref.set(qn("w:id"), cid)
    ref_run.append(comment_ref)

    # 插入：Start 在第一个 run 前，End 和 Reference 在段落末尾
    first_run = para_elem.find(qn("w:r"))
    if first_run is not None:
        para_elem.insert(list(para_elem).index(first_run), range_start)
    else:
        para_elem.append(range_start)
    para_elem.append(range_end)
    para_elem.append(ref_run)


def inject_annotations(
    doc, annotations, author="储颖", initials="储颖", min_start_index=None
):
    """
    向文档注入引用批注。

    annotations: list[dict]
      每个字典：
        - text: str — 批注文本
        - target: dict
            type: "keyword" — keyword + start_from
            type: "paragraph_index" — index

    返回注入计划列表（用于后续 comments.xml 合成）
    """
    max_id = _get_max_comment_id(doc)
    next_id = max_id + 1
    injection_plan = []
    annotation_start_index = (
        _find_annotation_start_index(doc)
        if min_start_index is None
        else max(0, int(min_start_index))
    )

    for ann in annotations:
        target = ann["target"]
        para_elem = None

        if target["type"] == "keyword":
            start_from = max(target.get("start_from", 0), annotation_start_index)
            _, para_elem = _find_paragraph_by_keyword(
                doc, target["keyword"], start_from
            )
        elif target["type"] == "paragraph_index":
            idx = target["index"]
            if annotation_start_index <= idx < len(doc.paragraphs):
                para_elem = doc.paragraphs[idx]._element

        if para_elem is not None:
            cid = next_id
            next_id += 1
            date_str = ann.get("date", datetime.now().strftime("%Y-%m-%dT%H:%M:00Z"))
            _mark_paragraph_with_comment(para_elem, str(cid))
            injection_plan.append(
                {
                    "comment_id": cid,
                    "client_comment_id": ann.get("comment_id", ""),
                    "text": ann["text"],
                    "date": date_str,
                    "author": author,
                    "initials": initials,
                    "evidence_kind": ann.get("evidence_kind", "text"),
                }
            )

    return injection_plan


# ══════════════════════════════════════════════════════════
# 模块1：成绩评定填充
# ══════════════════════════════════════════════════════════


def _find_grading_cell(doc, table_index=1, cell_row=0, cell_col=0):
    """定位成绩评定所在的表格单元格"""
    if table_index < len(doc.tables):
        table = doc.tables[table_index]
        if cell_row < len(table.rows):
            row = table.rows[cell_row]
            if cell_col < len(row.cells):
                return row.cells[cell_col]
    return None


def _find_para_by_prefix(cell, prefix):
    """在单元格中按段落文本前缀查找段落索引"""
    for i, p in enumerate(cell.paragraphs):
        if p.text.strip().startswith(prefix):
            return i, p
    return -1, None


def inject_score(doc, scores, table_index=1, cell_row=0, cell_col=0, late_days=0):
    """
    在「成绩评定：」行填写分数。

    scores: list[int/float] — 各维度分数列表
        例如 [37, 3, 3, 5, 7, 5, 23]
    late_days: int — 迟交天数（每天扣1分，默认0）

    将生成 "37+3+3+5+7+5+23=83" 格式的文本。
    若 late_days > 0，则生成 "37+3+3+5+7+5+23=83-3(延期)=80" 格式。
    """
    cell = _find_grading_cell(doc, table_index, cell_row, cell_col)
    if cell is None:
        print("警告：未找到成绩评定表格单元格", file=sys.stderr)
        return False

    idx, para = _find_para_by_prefix(cell, "成绩评定")
    if para is None:
        # 尝试查找空的"成绩评定"段落
        for i, p in enumerate(cell.paragraphs):
            if "成绩评定" in p.text:
                idx, para = i, p
                break

    if para is None:
        print("警告：未找到「成绩评定：」段落", file=sys.stderr)
        return False

    # 清除段落中已有的 runs（保留段落属性）
    p_elem = para._element
    # 移除所有 w:r 子元素
    for run_elem in p_elem.findall(qn("w:r")):
        p_elem.remove(run_elem)

    # 获取已有段落的样式属性（用于复制到新run）
    pPr = p_elem.find(qn("w:pPr"))
    rPr_template = None
    if pPr is not None:
        rPr_template = pPr.find(qn("w:rPr"))

    # 创建"成绩评定： "的run
    run1 = OxmlElement("w:r")
    if rPr_template is not None:
        run1.append(copy.deepcopy(rPr_template))
    else:
        rPr1 = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:hint"), "eastAsia")
        rPr1.append(rFonts)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), "21")
        rPr1.append(szCs)
        run1.append(rPr1)
    t1 = OxmlElement("w:t")
    t1.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t1.text = "成绩评定： "
    run1.append(t1)
    p_elem.append(run1)

    # 创建分数公式的run
    total = sum(scores)
    formula = "+".join(str(int(s)) for s in scores) + "=" + str(int(total))
    # 迟交扣分：在基础分之后追加 "-N(延期)=最终分"
    if late_days > 0:
        final_score = max(0, int(total) - late_days)
        formula += f"-{late_days}(延期)={final_score}"

    run2 = OxmlElement("w:r")
    if rPr_template is not None:
        run2.append(copy.deepcopy(rPr_template))
    else:
        rPr2 = OxmlElement("w:rPr")
        rFonts2 = OxmlElement("w:rFonts")
        rFonts2.set(qn("w:hint"), "eastAsia")
        rPr2.append(rFonts2)
        szCs2 = OxmlElement("w:szCs")
        szCs2.set(qn("w:val"), "21")
        rPr2.append(szCs2)
        run2.append(rPr2)
    t2 = OxmlElement("w:t")
    t2.text = formula
    run2.append(t2)
    p_elem.append(run2)

    print(f"  成绩评定已填写: {formula}")
    return True


# ══════════════════════════════════════════════════════════
# 模块2：评语填写
# ══════════════════════════════════════════════════════════


def inject_comment_text(doc, comment_text, table_index=1, cell_row=0, cell_col=0):
    """
    在「评语：」行填写评语文本。

    comment_text: str — 评语内容（不含前缀"评语："）
    """
    cell = _find_grading_cell(doc, table_index, cell_row, cell_col)
    if cell is None:
        print("警告：未找到评语表格单元格", file=sys.stderr)
        return False

    idx, para = _find_para_by_prefix(cell, "评语")
    if para is None:
        # 查找包含"评语"的段落
        for i, p in enumerate(cell.paragraphs):
            if "评语" in p.text:
                idx, para = i, p
                break

    if para is None:
        print("警告：未找到「评语：」段落", file=sys.stderr)
        return False

    # 清除已有 runs
    p_elem = para._element
    for run_elem in p_elem.findall(qn("w:r")):
        p_elem.remove(run_elem)

    # 获取样式模板
    pPr = p_elem.find(qn("w:pPr"))
    rPr_template = None
    if pPr is not None:
        rPr_template = pPr.find(qn("w:rPr"))

    def _make_rPr():
        if rPr_template is not None:
            return copy.deepcopy(rPr_template)
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:hint"), "eastAsia")
        rPr.append(rFonts)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), "21")
        rPr.append(szCs)
        return rPr

    # "评语：" run
    run1 = OxmlElement("w:r")
    run1.append(_make_rPr())
    t1 = OxmlElement("w:t")
    t1.text = "评语："
    run1.append(t1)
    p_elem.append(run1)

    # 评语内容 run
    run2 = OxmlElement("w:r")
    run2.append(_make_rPr())
    t2 = OxmlElement("w:t")
    t2.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t2.text = comment_text
    run2.append(t2)
    p_elem.append(run2)

    print(f"  评语已填写: {comment_text[:60]}...")
    return True


def append_generic_grading_section(doc, criteria, scores, comment_text):
    """
    为没有固定“成绩评定/评语”表格的通用 Word 追加可编辑批改页。

    原文、图片和已有版式保持不变；批改结果作为独立附页写入，避免通用
    报告因为不符合某一所学校的模板而丢失分项成绩或教师评语。
    """
    if not scores and not comment_text:
        return False

    def set_east_asia_font(run, font_name="微软雅黑"):
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        r_fonts.set(qn("w:eastAsia"), font_name)

    doc.add_page_break()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("教师批改意见")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(23, 77, 64)
    title_run.font.name = "Arial"
    set_east_asia_font(title_run)

    note = doc.add_paragraph("LabTrace 证据化批改 · 可编辑 Word 交付")
    note.paragraph_format.space_after = Pt(10)
    for run in note.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(104, 117, 111)

    normalized_criteria = list(criteria or [])
    if not normalized_criteria:
        normalized_criteria = [
            {
                "name": f"评分维度 {index + 1}",
                "score": score,
                "max_score": "",
                "reason": "",
            }
            for index, score in enumerate(scores or [])
        ]

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    headers = ("评分维度", "得分", "满分", "证据化评语")
    for cell, header in zip(table.rows[0].cells, headers):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = header
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "174D40")
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
                set_east_asia_font(run)

    for criterion in normalized_criteria:
        cells = table.add_row().cells
        cells[0].text = str(criterion.get("name", "评分维度"))
        cells[1].text = f"{float(criterion.get('score', 0)):g}"
        max_score = criterion.get("max_score", "")
        cells[2].text = f"{float(max_score):g}" if max_score != "" else ""
        cells[3].text = str(criterion.get("reason", ""))[:500]
        for column, cell in enumerate(cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if column in {1, 2}
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    set_east_asia_font(run)

    total = sum(float(item.get("score", 0)) for item in normalized_criteria)
    maximum = sum(float(item.get("max_score", 0) or 0) for item in normalized_criteria)
    total_paragraph = doc.add_paragraph()
    total_paragraph.paragraph_format.space_before = Pt(10)
    total_run = total_paragraph.add_run(
        f"最终成绩：{total:g}" + (f" / {maximum:g}" if maximum else "")
    )
    total_run.bold = True
    total_run.font.size = Pt(12)
    total_run.font.color.rgb = RGBColor(23, 77, 64)
    set_east_asia_font(total_run)

    if comment_text:
        comment_paragraph = doc.add_paragraph()
        comment_paragraph.paragraph_format.space_before = Pt(6)
        label = comment_paragraph.add_run("教师评语：")
        label.bold = True
        body = comment_paragraph.add_run(str(comment_text))
        for run in (label, body):
            run.font.size = Pt(10.5)
            set_east_asia_font(run)
    return True


def append_evidence_reference_appendix(doc, references):
    """Append a research-style evidence index as the final editable Word page.

    ``references`` is deliberately presentation-oriented.  It is generated by
    the API from the validated grading trace, so the Word appendix and the web
    console use the same stable ``[n]`` citation sequence.
    """
    normalized = list(references or [])
    if not normalized:
        return False

    def set_east_asia_font(run, font_name="微软雅黑"):
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        r_fonts.set(qn("w:eastAsia"), font_name)

    doc.add_page_break()
    title = doc.add_paragraph()
    title_run = title.add_run("附录：证据引用索引（LabTrace）")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(23, 77, 64)
    set_east_asia_font(title_run)

    note = doc.add_paragraph(
        "评分理由与 Word 批注中的 [n] 引用均指向本表；证据编号保留解析器原始定位，便于复核。"
    )
    note.paragraph_format.space_after = Pt(9)
    for run in note.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(104, 117, 111)
        set_east_asia_font(run)

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    headers = ("引用", "证据编号", "Word 原文位置", "证据类型", "摘要 / 关联评分维度")
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "174D40")
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(255, 255, 255)
                set_east_asia_font(run)

    for reference in normalized:
        cells = table.add_row().cells
        linked = "、".join(str(item) for item in reference.get("linked_criteria", []))
        excerpt = str(reference.get("excerpt", "")).strip()
        cells[0].text = f"[{int(reference.get('reference_number', 0))}]"
        cells[1].text = str(reference.get("evidence_id", ""))
        cells[2].text = str(reference.get("location_label", ""))
        cells[3].text = str(reference.get("kind_label", ""))
        cells[4].text = excerpt[:420] + (f"\n关联：{linked}" if linked else "")
        for column, cell in enumerate(cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if column == 0
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    set_east_asia_font(run)
    return True


# ══════════════════════════════════════════════════════════
# 模块3：签名图片 + 日期注入
# ══════════════════════════════════════════════════════════


def inject_signature_and_date(
    doc,
    signature_image_path,
    date_str=None,
    table_index=1,
    cell_row=0,
    cell_col=0,
    sig_width_inches=1.1,
    sig_height_inches=0.62,
    locate_keyword="指导教师签字",
    locate_fallback_keywords=None,
):
    """
    在「指导教师签字：」后插入签名图片，并在日期段落填写日期。

    signature_image_path: str — 签名图片路径
    date_str: str — 日期字符串，格式 "YYYY年 MM 月 DD 日"
                    如果为 None，使用当前日期
    locate_keyword: str — 主定位辅助词，用于在文档中搜索签名位置
    locate_fallback_keywords: list[str] — 备用定位辅助词列表
    """
    cell = _find_grading_cell(doc, table_index, cell_row, cell_col)
    if cell is None:
        print("警告：未找到签名表格单元格", file=sys.stderr)
        return False

    # 构建搜索关键词列表：主关键词在前，备用关键词在后
    keywords_to_try = [locate_keyword]
    if locate_fallback_keywords:
        keywords_to_try.extend(locate_fallback_keywords)

    sig_idx, sig_para = None, None
    for kw in keywords_to_try:
        # 先尝试精确前缀匹配
        idx, para = _find_para_by_prefix(cell, kw)
        if para is not None:
            sig_idx, sig_para = idx, para
            break
        # 再尝试包含匹配
        for i, p in enumerate(cell.paragraphs):
            if kw in p.text:
                sig_idx, sig_para = i, p
                break
        if sig_para is not None:
            break

    if sig_para is None:
        tried = ", ".join(f"「{kw}」" for kw in keywords_to_try)
        print(f"警告：未找到签名定位段落（已尝试: {tried}）", file=sys.stderr)
        return False

    # 检查是否已有签名图片
    has_existing_sig = False
    for run in sig_para.runs:
        drawings = run._element.findall(".//" + qn("w:drawing"))
        if drawings:
            has_existing_sig = True
            break

    if not has_existing_sig:
        # 插入签名图片
        if not os.path.exists(signature_image_path):
            print(f"警告：签名图片不存在: {signature_image_path}", file=sys.stderr)
            return False

        # 在签名段落末尾添加一个新的 run with inline image
        run = sig_para.add_run()
        run.add_picture(
            signature_image_path,
            width=Inches(sig_width_inches),
            height=Inches(sig_height_inches),
        )
        print(f"  签名图片已插入: {signature_image_path}")
    else:
        print("  签名图片已存在，跳过插入")

    # 处理日期
    if date_str is None:
        now = datetime.now()
        date_str = f"{now.year}年 {now.month:02d} 月 {now.day:02d} 日"

    # 定位日期段落（签名段落后第2个段落，即 sig_idx + 2）
    date_para = None
    # 策略1：查找包含"年"和"月"和"日"的段落
    for i, p in enumerate(cell.paragraphs):
        text = p.text.strip()
        if "年" in text and "月" in text and "日" in text and i > sig_idx:
            date_para = p
            break

    # 策略2：如果没找到，在 sig_idx + 2 位置
    if date_para is None and sig_idx + 2 < len(cell.paragraphs):
        date_para = cell.paragraphs[sig_idx + 2]

    if date_para is not None:
        # 清除已有 runs
        p_elem = date_para._element
        for run_elem in p_elem.findall(qn("w:r")):
            p_elem.remove(run_elem)

        # 获取样式
        pPr = p_elem.find(qn("w:pPr"))
        rPr_template = None
        if pPr is not None:
            rPr_template = pPr.find(qn("w:rPr"))

        def _make_rPr():
            if rPr_template is not None:
                return copy.deepcopy(rPr_template)
            rPr = OxmlElement("w:rPr")
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:hint"), "eastAsia")
            rPr.append(rFonts)
            szCs = OxmlElement("w:szCs")
            szCs.set(qn("w:val"), "21")
            rPr.append(szCs)
            return rPr

        # 空格缩进 + 日期
        run_space = OxmlElement("w:r")
        run_space.append(_make_rPr())
        t_space = OxmlElement("w:t")
        t_space.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t_space.text = "                                                         "
        run_space.append(t_space)
        p_elem.append(run_space)

        # 日期拆分为多个 run 以匹配原始格式
        # 解析日期字符串
        m = re.match(r"(\d{4})年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日", date_str)
        if m:
            year, month, day = m.group(1), m.group(2).zfill(2), m.group(3).zfill(2)
            parts = [year, f"年 ", f"{month} ", f"月 ", f"{day} ", f"日"]
        else:
            parts = [date_str]

        for part_text in parts:
            run_d = OxmlElement("w:r")
            run_d.append(_make_rPr())
            t_d = OxmlElement("w:t")
            t_d.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t_d.text = part_text
            run_d.append(t_d)
            p_elem.append(run_d)

        print(f"  日期已填写: {date_str}")
    else:
        print("  警告：未找到日期段落", file=sys.stderr)

    return True


# ══════════════════════════════════════════════════════════
# 模块4：Comments XML 合成（zipfile 级操作）
# ══════════════════════════════════════════════════════════


def _synthesize_comments_xml(doc_buffer, injection_plan):
    """
    在 docx zip 包中合成 comments.xml。

    doc_buffer: BytesIO — 已保存的 docx 文件 buffer
    injection_plan: list[dict] — 批注注入计划

    返回 BytesIO — 合成后的 docx buffer
    """
    if not injection_plan:
        return doc_buffer

    doc_buffer.seek(0)
    out_buf = BytesIO()

    with zipfile.ZipFile(doc_buffer, "r") as zin:
        has_comments = "word/comments.xml" in zin.namelist()

        with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                data = zin.read(item)

                if item == "word/comments.xml":
                    # 修改已有的 comments.xml
                    root = etree.fromstring(data)
                    for plan in injection_plan:
                        comment_elem = _create_comment_xml(
                            plan["comment_id"],
                            plan["author"],
                            plan["date"],
                            plan["initials"],
                            plan["text"],
                        )
                        root.append(comment_elem)
                    data = etree.tostring(
                        root, xml_declaration=True, encoding="UTF-8", standalone=True
                    )

                elif item == "[Content_Types].xml" and not has_comments:
                    ct_root = etree.fromstring(data)
                    # 添加 comments content type
                    ns = ct_root.nsmap.get(None, "")
                    override = etree.SubElement(
                        ct_root, "{%s}Override" % ns if ns else "Override"
                    )
                    override.set("PartName", "/word/comments.xml")
                    override.set("ContentType", COMMENTS_CT)
                    data = etree.tostring(
                        ct_root, xml_declaration=True, encoding="UTF-8", standalone=True
                    )

                elif item == "word/_rels/document.xml.rels" and not has_comments:
                    rels_root = etree.fromstring(data)
                    ns = rels_root.nsmap.get(None, "")
                    # 找最大 rId
                    max_rid = 0
                    for rel in rels_root:
                        rid = rel.get("Id", "")
                        if rid.startswith("rId"):
                            try:
                                max_rid = max(max_rid, int(rid[3:]))
                            except ValueError:
                                pass
                    new_rel = etree.SubElement(
                        rels_root, "{%s}Relationship" % ns if ns else "Relationship"
                    )
                    new_rel.set("Id", f"rId{max_rid + 1}")
                    new_rel.set("Type", COMMENTS_REL_TYPE)
                    new_rel.set("Target", "comments.xml")
                    data = etree.tostring(
                        rels_root,
                        xml_declaration=True,
                        encoding="UTF-8",
                        standalone=True,
                    )

                zout.writestr(item, data)

            if not has_comments:
                # 创建全新的 comments.xml
                NSMAP = {
                    "w": W_NS,
                    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                }
                comments_root = etree.Element(qn("w:comments"), nsmap=NSMAP)
                for plan in injection_plan:
                    comment_elem = _create_comment_xml(
                        plan["comment_id"],
                        plan["author"],
                        plan["date"],
                        plan["initials"],
                        plan["text"],
                    )
                    comments_root.append(comment_elem)
                comments_data = etree.tostring(
                    comments_root,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                )
                zout.writestr("word/comments.xml", comments_data)

    out_buf.seek(0)
    return out_buf


# ══════════════════════════════════════════════════════════
# 集成入口
# ══════════════════════════════════════════════════════════


def inject_all(input_path, output_path, config):
    """
    一站式注入全部批改内容。

    config: dict — 批改配置，结构如下：
    {
        "annotations": [                    # 可选：引用批注列表
            {
                "text": "批注内容",
                "target": {
                    "type": "keyword",      # "keyword" 或 "paragraph_index"
                    "keyword": "关键词",
                    "start_from": 0         # 可选
                }
            }
        ],
        "scores": [37, 3, 3, 5, 7, 5, 23], # 可选：各维度分数
        "comment": "评语内容",               # 可选：评语文本（不含"评语："前缀）
        "signature": {                       # 可选：签名配置
            "image_path": "path/to/sig.png",
            "date": "2025年 04 月 19 日",    # 可选，默认当前日期
            "width_inches": 1.1,             # 可选
            "height_inches": 0.62            # 可选
        },
        "author": "储颖",                    # 可选：批注作者
        "table_index": 1,                    # 可选：成绩表格索引
        "cell_row": 0,                       # 可选
        "cell_col": 0                        # 可选
    }
    """
    doc = Document(input_path)
    author = config.get("author", "储颖")
    initials = config.get("initials", author)
    table_idx = config.get("table_index", 1)
    cell_row = config.get("cell_row", 0)
    cell_col = config.get("cell_col", 0)

    results = {
        "annotations_count": 0,
        "injected_annotation_ids": [],
        "image_annotations_count": 0,
        "annotation_start_index": None,
        "score_injected": False,
        "comment_injected": False,
        "generic_section_appended": False,
        "evidence_appendix_appended": False,
        "delivery_mode": "comments_only",
        "signature_injected": False,
    }
    injection_plan = []

    # 0. 引用批注
    annotations = config.get("annotations", [])
    if annotations:
        annotation_start_index = _find_annotation_start_index(doc)
        results["annotation_start_index"] = annotation_start_index
        injection_plan = inject_annotations(
            doc,
            annotations,
            author,
            initials,
            min_start_index=annotation_start_index,
        )
        results["annotations_count"] = len(injection_plan)
        results["injected_annotation_ids"] = [
            item["client_comment_id"]
            for item in injection_plan
            if item.get("client_comment_id")
        ]
        results["image_annotations_count"] = sum(
            1 for item in injection_plan if item.get("evidence_kind") == "image"
        )
        print(
            f"[0/3] 批注注入: {len(injection_plan)} 条（从段落 {annotation_start_index} 开始定位）"
        )

    # 1. 成绩评定
    scores = config.get("scores", None)
    late_days = config.get("late_days", 0)
    if scores:
        results["score_injected"] = inject_score(
            doc, scores, table_idx, cell_row, cell_col, late_days=late_days
        )
        total = sum(scores)
        formula = "+".join(str(int(s)) for s in scores) + "=" + str(int(total))
        if late_days > 0:
            final_score = max(0, int(total) - late_days)
            formula += f"-{late_days}(延期)={final_score}"
        print(f"[1/3] 成绩评定: {formula}")

    # 2. 评语
    comment = config.get("comment", None)
    if comment:
        results["comment_injected"] = inject_comment_text(
            doc, comment, table_idx, cell_row, cell_col
        )
        print(f"[2/3] 评语: {comment[:60]}...")

    # 通用 Word 不一定包含某所学校的固定评分表。只要模板字段有任一项
    # 无法写入，就追加一页结构化、可编辑的教师批改意见，保证任务闭环。
    if (scores and not results["score_injected"]) or (
        comment and not results["comment_injected"]
    ):
        results["generic_section_appended"] = append_generic_grading_section(
            doc,
            config.get("criteria", []),
            scores or [],
            comment or "",
        )
        if results["generic_section_appended"]:
            results["score_injected"] = bool(scores)
            results["comment_injected"] = bool(comment)
            results["delivery_mode"] = "generic_appendix"
    elif results["score_injected"] or results["comment_injected"]:
        results["delivery_mode"] = "template_fields"

    # 3. 签名 + 日期
    sig_config = config.get("signature", None)
    if sig_config:
        sig_path = sig_config.get("image_path", "")
        sig_date = sig_config.get("date", None)
        sig_w = sig_config.get("width_inches", 1.1)
        sig_h = sig_config.get("height_inches", 0.62)
        sig_keyword = sig_config.get("locate_keyword", "指导教师签字")
        sig_fallback = sig_config.get("locate_fallback_keywords", None)
        results["signature_injected"] = inject_signature_and_date(
            doc,
            sig_path,
            sig_date,
            table_idx,
            cell_row,
            cell_col,
            sig_w,
            sig_h,
            locate_keyword=sig_keyword,
            locate_fallback_keywords=sig_fallback,
        )
        print(f"[3/3] 签名+日期: {sig_date or '当前日期'}")

    # 4. 科研式证据引用附录必须位于整份交付物末尾，使网页中的 [n]
    # 引用、Word 原生批注和内部 p-/t-/i- 定位符可以被教师交叉核验。
    references = config.get("evidence_appendix", [])
    if references:
        results["evidence_appendix_appended"] = append_evidence_reference_appendix(
            doc, references
        )

    # 保存到 buffer
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    # 如果有批注需要合成 comments.xml
    if injection_plan:
        buf = _synthesize_comments_xml(buf, injection_plan)

    # 写入输出文件
    buf.seek(0)
    with open(output_path, "wb") as f:
        f.write(buf.read())

    print(f"\n批改完成，输出: {output_path}")
    return results


# ── CLI ──────────────────────────────────────────────────
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="实验报告批改内容注入器（完整集成版）")
    parser.add_argument("--input", "-i", required=True, help="输入 DOCX 文件路径")
    parser.add_argument("--output", "-o", required=True, help="输出 DOCX 文件路径")
    parser.add_argument("--config", "-c", required=True, help="批改配置 JSON 文件路径")
    args = parser.parse_args()

    with open(args.config, "r", encoding="utf-8") as f:
        config = json.load(f)

    results = inject_all(args.input, args.output, config)
    print(f"\n结果汇总: {json.dumps(results, ensure_ascii=False, indent=2)}")
