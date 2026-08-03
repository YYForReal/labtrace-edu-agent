#!/usr/bin/env python3
"""
详细评价注入脚本
将优点、不足、建议、警告以颜色编码的方式注入到文档中
"""

import sys
import json
import argparse

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("请安装 python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def inject_comments(doc, grading_result):
    """
    在文档末尾注入详细评价

    Args:
        doc: python-docx Document 对象
        grading_result: 评分结果字典
    """
    doc.add_paragraph()

    # 标题
    comment_title = doc.add_paragraph()
    title_run = comment_title.add_run("【详细评价】")
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    # 详细分析
    analysis = grading_result.get("detailed_analysis", "")
    if analysis:
        doc.add_paragraph()
        analysis_para = doc.add_paragraph(analysis)
        for run in analysis_para.runs:
            run.font.size = Pt(11)

    # 优点
    strengths = grading_result.get("strengths", [])
    if strengths:
        doc.add_paragraph()
        s_para = doc.add_paragraph()
        s_run = s_para.add_run("✓ 优点：")
        s_run.font.bold = True
        s_run.font.color.rgb = RGBColor(0, 128, 0)
        s_run.font.size = Pt(12)

        for item in strengths:
            p = doc.add_paragraph(f"  • {item}")
            for run in p.runs:
                run.font.size = Pt(11)

    # 不足
    weaknesses = grading_result.get("weaknesses", [])
    if weaknesses:
        doc.add_paragraph()
        w_para = doc.add_paragraph()
        w_run = w_para.add_run("⚠ 需要改进：")
        w_run.font.bold = True
        w_run.font.color.rgb = RGBColor(204, 102, 0)
        w_run.font.size = Pt(12)

        for item in weaknesses:
            p = doc.add_paragraph(f"  • {item}")
            for run in p.runs:
                run.font.size = Pt(11)

    # 建议
    suggestions = grading_result.get("suggestions", [])
    if suggestions:
        doc.add_paragraph()
        sg_para = doc.add_paragraph()
        sg_run = sg_para.add_run("→ 改进建议：")
        sg_run.font.bold = True
        sg_run.font.color.rgb = RGBColor(0, 51, 153)
        sg_run.font.size = Pt(12)

        for item in suggestions:
            p = doc.add_paragraph(f"  • {item}")
            for run in p.runs:
                run.font.size = Pt(11)

    # 警告
    warnings = grading_result.get("warnings", [])
    if warnings:
        doc.add_paragraph()
        wr_para = doc.add_paragraph()
        wr_run = wr_para.add_run("⚠ 警告：")
        wr_run.font.bold = True
        wr_run.font.color.rgb = RGBColor(204, 0, 0)
        wr_run.font.size = Pt(12)

        for item in warnings:
            p = doc.add_paragraph(f"  • {item}")
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(204, 0, 0)


def main():
    parser = argparse.ArgumentParser(description="注入详细评价到 DOCX")
    parser.add_argument("--input", "-i", required=True, help="输入 DOCX 文件")
    parser.add_argument("--grading-result", "-g", required=True, help="评分结果 JSON")
    parser.add_argument("--output", "-o", required=True, help="输出 DOCX 文件")

    args = parser.parse_args()

    with open(args.grading_result, "r", encoding="utf-8") as f:
        grading_result = json.load(f)

    doc = Document(args.input)
    inject_comments(doc, grading_result)
    doc.save(args.output)

    print(f"详细评价已注入: {args.output}")


if __name__ == "__main__":
    main()
