#!/usr/bin/env python3
"""
DOCX 文档解析脚本
从 python-docx 提取文本、段落、表格、图片、结构和元数据。

通过标准 DOCX ZIP/XML 结构获取修订与批注诊断信息。
"""

import os
import sys
import json
import argparse
import re
import logging
import zipfile
from datetime import datetime

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
except ImportError:
    print("请安装 python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)

# ─── 核心解析函数 ─────────────────────────────────────────────────────


def parse_docx(file_path):
    """
    解析 DOCX 文件，返回结构化数据。

    解析流程：
    1. python-docx 提取段落 + 表格 + 图片 + 元数据
    2. 直接读取 DOCX ZIP 中的 XML，统计修订和批注

    Args:
        file_path: DOCX 文件路径

    Returns:
        dict: ParsedDocument 结构
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    doc = Document(file_path)

    # ── 1. 提取段落（同时检测段落内嵌图片，插入占位符） ────────
    paragraphs = []
    full_text_parts = []
    # image_placeholder_count: 记录已插入的 [IMAGE_N] 占位符数量
    image_placeholder_count = 0

    # XML 命名空间，用于检测段落中的嵌入图片
    _DRAWING_TAG = (
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
    )
    _BLIP_TAG = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"

    for para in doc.paragraphs:
        text = para.text.strip()

        style_name = para.style.name if para.style else "Normal"
        level = None

        # 检测标题级别
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1
        elif style_name == "Title":
            level = 0

        # 检测该段落是否包含嵌入图片（通过 XML 中的 <w:drawing> 元素）
        para_images = para._element.findall(f".//{_DRAWING_TAG}")
        img_count_in_para = 0
        for drawing in para_images:
            # 确认 drawing 中包含真正的图片引用（blip），排除其他 drawing 元素
            if drawing.findall(f".//{_BLIP_TAG}"):
                img_count_in_para += 1

        if img_count_in_para > 0:
            # 在该段落的文本后追加图片占位符
            placeholders = []
            for _ in range(img_count_in_para):
                image_placeholder_count += 1
                placeholders.append(f"[IMAGE_{image_placeholder_count}]")
            placeholder_text = " ".join(placeholders)

            if text:
                paragraphs.append({"text": text, "style": style_name, "level": level})
                full_text_parts.append(text)
                full_text_parts.append(placeholder_text)
            else:
                # 纯图片段落（无文字），只添加占位符
                full_text_parts.append(placeholder_text)
        else:
            if not text:
                continue
            paragraphs.append({"text": text, "style": style_name, "level": level})
            full_text_parts.append(text)

    # ── 2. 提取表格 ──────────────────────────────────────────────
    tables = _extract_tables(doc)
    # 将表格文本也加入 full_text，这样 LLM 能看到完整内容
    for tbl in tables:
        table_text = _table_to_text(tbl)
        if table_text.strip():
            full_text_parts.append(table_text)

    full_text = "\n".join(full_text_parts)

    # ── 3. 提取图片信息 ──────────────────────────────────────────
    images = _extract_images(doc, paragraphs)

    # ── 4. 构建文档结构树 ────────────────────────────────────────
    structure = _build_structure(paragraphs)

    # ── 5. 提取元数据 ────────────────────────────────────────────
    metadata = _extract_metadata(doc, full_text, images, tables)

    # ── 6. 提取学生信息 ──────────────────────────────────────────
    student_info = _extract_student_info_from_filename(file_path)

    # ── 7. 可选：docx skill XML 级深度解析 ───────────────────────
    xml_analysis = _try_xml_deep_parse(file_path)

    result = {
        "file_path": os.path.abspath(file_path),
        "file_type": "docx",
        "student_info": student_info,
        "full_text": full_text,
        "paragraphs": paragraphs,
        "tables": tables,
        "images": images,
        "image_placeholder_count": image_placeholder_count,
        "structure": structure,
        "metadata": metadata,
    }

    if xml_analysis:
        result["xml_analysis"] = xml_analysis

    return result


# ─── 表格解析 ────────────────────────────────────────────────────────


def _extract_tables(doc):
    """
    提取文档中所有表格的结构化数据。

    Returns:
        list[dict]: 每个表格包含 rows/cols/data/summary
    """
    tables = []

    for idx, table in enumerate(doc.tables):
        rows_data = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                # 提取单元格文本（含段落换行）
                cell_text = "\n".join(
                    p.text.strip() for p in cell.paragraphs if p.text.strip()
                )
                cells.append(cell_text)
            rows_data.append(cells)

        # 生成表格摘要（便于 LLM 快速理解表格用途）
        summary = _summarize_table(rows_data, idx)

        tables.append(
            {
                "index": idx,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "data": rows_data,
                "summary": summary,
            }
        )

    return tables


def _summarize_table(rows_data, table_index):
    """根据表格内容生成简短摘要"""
    if not rows_data:
        return f"空表格 #{table_index}"

    # 获取第一行的关键词
    first_row_text = " ".join(rows_data[0]).strip()[:200]

    # 检测常见表格类型
    keywords_map = {
        "课程": "课程信息表",
        "成绩": "成绩表",
        "评分": "评分表",
        "批阅": "教师批阅意见表",
        "指导教师": "教师批阅意见表",
        "备注": "备注表",
        "实验": "实验信息表",
    }

    for keyword, label in keywords_map.items():
        if keyword in first_row_text:
            return label

    return f"表格 #{table_index} ({len(rows_data)}行)"


def _table_to_text(table_dict):
    """将表格结构转为可读的纯文本，方便拼入 full_text"""
    lines = [f"\n[{table_dict['summary']}]"]
    for row in table_dict["data"]:
        row_text = " | ".join(cell.replace("\n", " ") for cell in row)
        if row_text.strip():
            lines.append(row_text)
    return "\n".join(lines)


# ─── XML 深度解析 ──────────────────────────────────────────────────


def _try_xml_deep_parse(file_path):
    """
    直接读取 DOCX ZIP 包，提取 tracked changes 与批注数量。

    Returns:
        dict | None
    """
    try:
        analysis = {}
        with zipfile.ZipFile(file_path) as package:
            names = set(package.namelist())
            content = ""
            if "word/document.xml" in names:
                content = package.read("word/document.xml").decode(
                    "utf-8", errors="replace"
                )
            analysis["has_tracked_changes"] = (
                "<w:ins " in content or "<w:del " in content
            )
            analysis["tracked_insertions"] = content.count("<w:ins ")
            analysis["tracked_deletions"] = content.count("<w:del ")

            if "word/comments.xml" in names:
                comments = package.read("word/comments.xml").decode(
                    "utf-8", errors="replace"
                )
                analysis["comment_count"] = comments.count("<w:comment ")

        return analysis if analysis else None

    except Exception as exc:
        logger.debug("XML 深度解析失败: %s", exc)
        return None


def _extract_images(doc, paragraphs):
    """
    按文档阅读顺序提取所有嵌入图片的信息和二进制数据。

    重要：图片按段落中出现的顺序枚举，与 full_text 中 [IMAGE_N] 占位符一一对应。
    不使用 doc.part.rels 枚举（其顺序不保证与文档阅读顺序一致）。

    遍历逻辑：
      doc.paragraphs → 每个段落 XML 中查找 <w:drawing> → <a:blip r:embed="rIdXX">
      → 通过 rId 从 doc.part.rels 获取实际图片 Part → 提取 base64

    去重逻辑：
      对每张图片的二进制内容计算 SHA-256 hash，若与前面已出现的图片完全相同，
      标记 duplicate_of = 首次出现的 image index，跳过 Vision 分析以节省 API 调用。

    Returns:
        list[dict]: 每个图片包含 index, description, context, base64, media_type, size_bytes,
                    以及可选的 duplicate_of (int) 和 content_hash (str)
    """
    import base64 as _b64
    import hashlib

    images = []
    image_index = 0

    # 去重用：content_hash → 首次出现的 image_index
    _seen_hashes: dict[str, int] = {}

    # XML 命名空间常量
    _W_DRAWING = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
    _A_BLIP = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    _R_EMBED = (
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
    )

    # 常见图片 MIME 类型映射
    _EXT_MIME = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".bmp": "image/bmp",
        ".tiff": "image/tiff",
        ".tif": "image/tiff",
        ".emf": "image/emf",
        ".wmf": "image/wmf",
    }

    # 建立 doc.paragraphs → paragraphs 列表的索引映射
    # paragraphs 列表跳过了空段落，所以需要按文本内容匹配
    _para_text_to_idx = {}
    for pi, p in enumerate(paragraphs):
        text = p["text"]
        if text not in _para_text_to_idx:
            _para_text_to_idx[text] = pi

    # 追踪最近匹配的 paragraphs 索引，用于图片上下文定位
    _last_para_idx = 0

    for docx_paragraph_index, para in enumerate(doc.paragraphs):
        # 尝试将当前 doc 段落映射到 paragraphs 列表的索引
        para_text = para.text.strip()
        if para_text and para_text in _para_text_to_idx:
            _last_para_idx = _para_text_to_idx[para_text]

        drawings = para._element.findall(f".//{_W_DRAWING}")
        for drawing in drawings:
            blips = drawing.findall(f".//{_A_BLIP}")
            for blip in blips:
                embed_rid = blip.get(_R_EMBED)
                if not embed_rid:
                    continue

                context = _get_image_context(
                    paragraphs, image_index, para_index=_last_para_idx
                )

                image_info = {
                    "index": image_index,
                    "description": f"嵌入图片 {image_index + 1}",
                    "context": context,
                    # Keep both coordinates: ``paragraph_index`` addresses the
                    # compact parsed paragraph catalog, while
                    # ``docx_paragraph_index`` addresses the original Word
                    # document and can therefore anchor a native comment back
                    # to the image-containing paragraph.
                    "paragraph_index": _last_para_idx,
                    "docx_paragraph_index": docx_paragraph_index,
                }

                # 通过 rId 从文档 rels 获取图片 Part
                try:
                    rel = doc.part.rels.get(embed_rid)
                    if rel is None:
                        image_info["extraction_error"] = (
                            f"未找到 relationship: {embed_rid}"
                        )
                        images.append(image_info)
                        image_index += 1
                        continue

                    image_part = rel.target_part
                    image_blob = image_part.blob
                    image_name = image_part.partname  # 如 /word/media/image1.png

                    ext = os.path.splitext(str(image_name))[1].lower()
                    media_type = _EXT_MIME.get(ext, "image/png")

                    # 跳过矢量格式（EMF/WMF），LLM 无法理解
                    if ext in (".emf", ".wmf"):
                        image_info["skipped"] = True
                        image_info["skip_reason"] = f"矢量格式 {ext} 不适合多模态分析"
                    else:
                        image_info["base64"] = _b64.b64encode(image_blob).decode(
                            "ascii"
                        )
                        image_info["media_type"] = media_type
                        image_info["size_bytes"] = len(image_blob)

                        # ── 去重检测：基于图片内容 SHA-256 hash ──
                        content_hash = hashlib.sha256(image_blob).hexdigest()
                        image_info["content_hash"] = content_hash

                        if content_hash in _seen_hashes:
                            first_idx = _seen_hashes[content_hash]
                            image_info["duplicate_of"] = first_idx
                            logger.debug(
                                "图片 %d 与图片 %d 内容完全相同 (hash=%s...)",
                                image_index + 1,
                                first_idx + 1,
                                content_hash[:12],
                            )
                        else:
                            _seen_hashes[content_hash] = image_index

                except Exception as exc:
                    logger.debug(
                        "提取图片 %d 失败 (rId=%s): %s", image_index, embed_rid, exc
                    )
                    image_info["extraction_error"] = str(exc)

                images.append(image_info)
                image_index += 1

    # 统计去重结果
    dup_count = sum(1 for img in images if "duplicate_of" in img)
    if dup_count > 0:
        logger.info(
            "图片去重: %d/%d 张为重复图片（将复用首次分析结果）", dup_count, len(images)
        )

    return images


def _get_image_context(paragraphs, image_index, para_index=None):
    """
    获取图片周围的文本上下文。

    Args:
        paragraphs: 段落列表
        image_index: 图片序号（0-based）
        para_index: 图片所在段落在 paragraphs 中的索引（精确定位）；
                    若为 None 则用 image_index 作为近似位置。

    策略：取图片所在段落的前 2 段 + 当前段 + 后 2 段，
    每段最多 200 字符，提供足够上下文帮助 Vision 模型理解图片含义。
    """
    if not paragraphs:
        return ""

    idx = (
        para_index if para_index is not None else min(image_index, len(paragraphs) - 1)
    )
    context_parts = []

    # 向前取 2 段 + 当前段 + 向后取 2 段
    start = max(0, idx - 2)
    end = min(len(paragraphs), idx + 3)

    for i in range(start, end):
        text = paragraphs[i]["text"][:200]
        if not text:
            continue
        # 标注当前段落（图片所在段落）
        if i == idx:
            context_parts.append(f"[当前段落] {text}")
        else:
            context_parts.append(text)

    return " | ".join(context_parts)


def _build_structure(paragraphs):
    """从段落列表构建文档结构树"""
    sections = []
    current_section = None
    content_length = 0

    for para in paragraphs:
        if para["level"] is not None and para["level"] >= 1:
            # 保存上一个 section
            if current_section:
                current_section["content_length"] = content_length
                sections.append(current_section)

            current_section = {
                "title": para["text"],
                "level": para["level"],
                "content_length": 0,
            }
            content_length = 0
        else:
            content_length += len(para["text"])

    # 保存最后一个 section
    if current_section:
        current_section["content_length"] = content_length
        sections.append(current_section)

    return {"sections": sections}


def _extract_metadata(doc, full_text, images, tables=None):
    """提取文档元数据"""
    core_props = doc.core_properties

    metadata = {
        "author": core_props.author or "",
        "created": str(core_props.created) if core_props.created else "",
        "modified": str(core_props.modified) if core_props.modified else "",
        "word_count": len(full_text),
        "char_count": len(full_text.replace(" ", "").replace("\n", "")),
        "paragraph_count": len(doc.paragraphs),
        "image_count": len(images),
        "table_count": len(tables) if tables else len(doc.tables),
    }

    return metadata


def _extract_student_info_from_filename(file_path):
    """从文件名提取学生信息"""
    basename = os.path.basename(file_path)
    name_without_ext = os.path.splitext(basename)[0]

    # 模式1: 10位学号 + 中文姓名（无分隔符）
    match = re.search(r"(\d{10})([\u4e00-\u9fff]{2,4})", name_without_ext)
    if match:
        return {"student_id": match.group(1), "name": match.group(2)}

    # 模式2: 学号_姓名 或 学号 姓名
    match = re.search(r"(\d{8,12})[_\s\-]([\u4e00-\u9fff]{2,4})", name_without_ext)
    if match:
        return {"student_id": match.group(1), "name": match.group(2)}

    # 模式3: 仅学号
    match = re.search(r"(\d{10})", name_without_ext)
    if match:
        return {"student_id": match.group(1), "name": ""}

    return None


def main():
    parser = argparse.ArgumentParser(description="解析 DOCX 文档")
    parser.add_argument("--input", "-i", required=True, help="输入 DOCX 文件路径")
    parser.add_argument("--output", "-o", help="输出 JSON 文件路径")

    args = parser.parse_args()

    try:
        result = parse_docx(args.input)

        output_json = json.dumps(result, ensure_ascii=False, indent=2)

        if args.output:
            with open(args.output, "w", encoding="utf-8") as f:
                f.write(output_json)
            print(f"解析完成，结果保存到: {args.output}")
        else:
            print(output_json)

    except Exception as e:
        print(f"错误: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
