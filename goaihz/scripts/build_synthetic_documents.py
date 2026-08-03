"""Build the public, fully synthetic DOCX fixtures used by the live demo.

The document deliberately contains a realistic but imperfect analysis section so
the demo's low-confidence human-review branch is exercised. It contains no real
student identity, institution record, contact information, or source submission.
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "goaihz" / "data" / "synthetic"
INPUT_PATH = OUTPUT_DIR / "demo-student-001_实验报告.docx"
GRADED_PATH = OUTPUT_DIR / "demo-student-001_实验报告_批改示例.docx"

NAVY = "173E59"
TEAL = "168A82"
ORANGE = "D66443"
LIGHT = "EEF3F2"
MUTED = "5D6B73"
DOC_FONT = "FandolHei"


def _force_style_font(style, *, size: int | None = None) -> None:
    style.font.name = DOC_FONT
    if size is not None:
        style.font.size = Pt(size)
    properties = style._element.get_or_add_rPr()
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), DOC_FONT)
    for attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        qualified = qn(f"w:{attribute}")
        if qualified in fonts.attrib:
            del fonts.attrib[qualified]


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_width(table, width_twips: int = 9360) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(width_twips))
    width.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell in row.cells:
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("格物智评公开演示 · 匿名合成用例    ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def _add_header_footer(document: Document) -> None:
    for section in document.sections:
        header = section.header.paragraphs[0]
        header.text = "LABTRACE / SYNTHETIC REPORT"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run.font.name = DOC_FONT
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MUTED)
        footer = section.footer.paragraphs[0]
        _page_number(footer)
        for run in footer.runs:
            run.font.name = DOC_FONT
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MUTED)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    _force_style_font(normal, size=11)
    normal.font.color.rgb = RGBColor.from_string("26343B")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = document.styles["Title"]
    _force_style_font(title, size=30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)

    for style_name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)):
        style = document.styles[style_name]
        _force_style_font(style, size=size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    defaults = document.styles.element.find(qn("w:docDefaults"))
    run_defaults = defaults.find(qn("w:rPrDefault")).find(qn("w:rPr"))
    fonts = run_defaults.find(qn("w:rFonts"))
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), DOC_FONT)
    for attribute in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        qualified = qn(f"w:{attribute}")
        if qualified in fonts.attrib:
            del fonts.attrib[qualified]
    run_defaults.find(qn("w:lang")).set(qn("w:eastAsia"), "zh-CN")


def _add_section_title(document: Document, number: str, title: str) -> None:
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.add_run(f"{number}  ").font.color.rgb = RGBColor.from_string(ORANGE)
    paragraph.add_run(title)


def _add_callout(document: Document, label: str, body: str) -> None:
    table = document.add_table(rows=1, cols=2)
    _set_table_width(table)
    table.columns[0].width = Cm(2.2)
    table.columns[1].width = Cm(13.6)
    left, right = table.rows[0].cells
    _set_cell_shading(left, TEAL)
    _set_cell_shading(right, LIGHT)
    p_left = left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_left.add_run(label)
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    right.paragraphs[0].add_run(body)


def _add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    p.add_run("SYNTHETIC / 公开演示材料").font.color.rgb = RGBColor.from_string(ORANGE)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(52)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("高校通用实验报告")
    run.font.name = DOC_FONT
    run.font.size = Pt(17)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)

    title = document.add_paragraph(style="Title")
    title.add_run("温度传感器标定与\n线性误差分析")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(36)
    run = subtitle.add_run("面向数据采集类实验的匿名合成样例")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    metadata = document.add_table(rows=5, cols=2)
    _set_table_width(metadata)
    metadata.columns[0].width = Cm(4)
    metadata.columns[1].width = Cm(11.8)
    rows = [
        ("课程", "实验方法与数据分析（演示课程）"),
        ("用例编号", "DEMO-STUDENT-001"),
        ("提交身份", "匿名学生 A（人工构造）"),
        ("实验日期", "2026 年 7 月 18 日"),
        ("报告版本", "v1.0 / GOAI 公开演示"),
    ]
    for index, (label, value) in enumerate(rows):
        left, right = metadata.rows[index].cells
        _set_cell_shading(left, LIGHT)
        left.paragraphs[0].add_run(label).bold = True
        right.paragraphs[0].add_run(value)

    document.add_paragraph()
    _add_callout(
        document,
        "数据声明",
        "本报告的身份、过程、表格、图像和结果均为人工构造的演示数据，仅用于验证批改 Agent；不对应任何真实学生、教师或院校记录。",
    )
    document.add_page_break()


def _add_method_page(document: Document) -> None:
    _add_section_title(document, "01", "实验目标与原理")
    document.add_paragraph(
        "实验目标：使用五个温度点对模拟电阻式温度传感器进行标定，建立输出电压与温度之间的线性关系，并评价重复测量的稳定性。"
    )
    document.add_paragraph(
        "实验原理：在给定工作区间内，传感器输出可近似写为 V = aT + b。通过最小二乘法估计斜率 a 和截距 b，再用预测值与观测值的差计算残差。若残差相对满量程较小，可认为该线性模型满足本次演示要求。"
    )

    _add_section_title(document, "02", "实验环境与方法")
    document.add_paragraph(
        "实验环境：模拟恒温槽 20–60 ℃；16 位数据采集模块；采样频率 10 Hz；每个温度点稳定 60 s 后记录 30 s 均值；软件环境为 Python 3.11 与常规数据分析库。"
    )
    document.add_paragraph(
        "实验方法：先在 20 ℃完成零点检查；随后依次设定 20、30、40、50、60 ℃。每个点重复测量三次，记录平均电压。完成全部采样后拟合线性模型，并比较各温度点的观测均值和模型预测值。"
    )

    table = document.add_table(rows=1, cols=4)
    _set_table_width(table)
    headers = ["阶段", "关键操作", "控制参数", "输出证据"]
    for cell, value in zip(table.rows[0].cells, headers):
        _set_cell_shading(cell, NAVY)
        run = cell.paragraphs[0].add_run(value)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    _set_repeat_table_header(table.rows[0])
    data = [
        ("预热", "传感器通电并等待稳定", "10 min", "零点记录"),
        ("设点", "设置恒温槽目标温度", "20–60 ℃", "温度读数"),
        ("采样", "稳定后记录三次均值", "10 Hz / 30 s", "电压数据"),
        ("拟合", "最小二乘线性回归", "V = aT + b", "参数与残差"),
    ]
    for row_values in data:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            cell.text = value
            _set_cell_margins(cell)

    document.add_paragraph()
    _add_callout(
        document,
        "复现提示",
        "本用例保留了环境、参数、重复次数与处理方法，但未记录恒温槽实际波动范围，用于触发“方法细节仍可补充”的评分反馈。",
    )
    document.add_page_break()


def _make_chart(path: Path) -> None:
    temperatures = [20, 30, 40, 50, 60]
    voltage = [0.82, 1.19, 1.61, 1.98, 2.43]
    fit = [0.801, 1.203, 1.605, 2.007, 2.409]
    width, height = 1296, 612
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_path = "/System/Library/Fonts/STHeiti Medium.ttc"
    font_regular = ImageFont.truetype(font_path, 23)
    font_small = ImageFont.truetype(font_path, 19)
    font_title = ImageFont.truetype(font_path, 34)
    left, top, right, bottom = 116, 112, 1220, 510
    draw.rounded_rectangle((55, 45, 1250, 565), radius=20, fill="#F7F9F8")
    draw.text(
        (88, 66), "温度—电压标定结果（合成数据）", font=font_title, fill="#173E59"
    )

    def point(t: float, value: float) -> tuple[int, int]:
        x = left + int((t - 20) / 40 * (right - left))
        y = bottom - int((value - 0.6) / 2.0 * (bottom - top))
        return x, y

    for tick in (0.8, 1.2, 1.6, 2.0, 2.4):
        y = point(20, tick)[1]
        draw.line((left, y, right, y), fill="#D9E4E1", width=2)
        draw.text((62, y - 12), f"{tick:.1f}", font=font_small, fill="#62716B")
    draw.line((left, top, left, bottom), fill="#77847F", width=2)
    draw.line((left, bottom, right, bottom), fill="#77847F", width=2)
    for temperature in temperatures:
        x = point(temperature, 0.6)[0]
        draw.line((x, bottom, x, bottom + 8), fill="#77847F", width=2)
        draw.text(
            (x - 15, bottom + 15), str(temperature), font=font_small, fill="#62716B"
        )
    draw.text((560, 548), "温度 / ℃", font=font_regular, fill="#4A5B54")
    draw.text((72, 118), "输出电压 / V", font=font_small, fill="#4A5B54")

    fit_points = [point(t, value) for t, value in zip(temperatures, fit)]
    draw.line(fit_points, fill="#168A82", width=6, joint="curve")
    for t, value in zip(temperatures, voltage):
        x, y = point(t, value)
        draw.ellipse(
            (x - 9, y - 9, x + 9, y + 9), fill="#D66443", outline="white", width=3
        )
    draw.line((930, 82, 978, 82), fill="#168A82", width=6)
    draw.text((990, 69), "线性拟合", font=font_small, fill="#4A5B54")
    draw.ellipse((1085, 74, 1101, 90), fill="#D66443")
    draw.text((1110, 69), "观测均值", font=font_small, fill="#4A5B54")
    image.save(path, "PNG")


def _add_results_page(document: Document, chart_path: Path) -> None:
    _add_section_title(document, "03", "实验数据与结果")
    document.add_paragraph(
        "实验数据：五个温度点均完成三次重复测量。表 1 列出三次结果、平均值与线性模型预测值，所有数值均为本项目人工构造。"
    )

    table = document.add_table(rows=1, cols=6)
    _set_table_width(table)
    headers = [
        "温度 / ℃",
        "第 1 次 / V",
        "第 2 次 / V",
        "第 3 次 / V",
        "均值 / V",
        "预测 / V",
    ]
    for cell, value in zip(table.rows[0].cells, headers):
        _set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    _set_repeat_table_header(table.rows[0])
    values = [
        ("20", "0.81", "0.82", "0.83", "0.82", "0.801"),
        ("30", "1.18", "1.20", "1.19", "1.19", "1.203"),
        ("40", "1.60", "1.62", "1.61", "1.61", "1.605"),
        ("50", "1.97", "1.99", "1.98", "1.98", "2.007"),
        ("60", "2.41", "2.44", "2.44", "2.43", "2.409"),
    ]
    for row_values in values:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            cell.text = value
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_margins(cell, top=60, bottom=60)

    caption = document.add_paragraph("表 1  温度传感器标定数据（匿名合成）")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(chart_path), width=Inches(6.35))
    caption = document.add_paragraph("图 1  温度—电压观测均值与线性拟合结果")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    document.add_paragraph(
        "运行结果：拟合得到 V = 0.0402T − 0.003，决定系数 R² = 0.9993。五个温度点的最大绝对残差为 0.027 V。"
    )
    document.add_page_break()


def _add_analysis_page(document: Document) -> None:
    _add_section_title(document, "04", "结果分析与验证")
    document.add_paragraph(
        "结果分析：输出电压随温度升高而近似线性增加，拟合曲线与观测均值接近。50 ℃点的观测值低于预测值，可能来自恒温槽波动或读数误差。整体来看，实验结果与预期一致。"
    )
    document.add_paragraph(
        "误差分析：可能的误差来源包括供电波动、传感器自热以及温度未完全稳定。本报告没有进一步计算标准差、置信区间或满量程误差，也未安排另一支标准传感器进行对照，因此对模型可靠性的判断仍然偏定性。"
    )

    _add_section_title(document, "05", "实验结论与反思")
    document.add_paragraph(
        "实验结论：本次实验完成了 20–60 ℃区间的五点标定，得到近似线性的温度—电压关系。重复数据差异较小，说明采集过程基本稳定。"
    )
    document.add_paragraph(
        "改进方向：后续应增加每个温度点的重复次数，报告标准差和置信区间；补充对照测量，并解释 50 ℃点偏差是否具有稳定性。结论还应逐项引用表 1 或图 1 中的具体数据。"
    )

    _add_section_title(document, "06", "教师批阅区")
    grading = document.add_table(rows=1, cols=1)
    _set_table_width(grading)
    cell = grading.cell(0, 0)
    _set_cell_shading(cell, "F7F9F8")
    cell.paragraphs[0].add_run("指导教师批阅意见：").bold = True
    cell.add_paragraph("成绩评定：")
    cell.add_paragraph("评语：")
    cell.add_paragraph("指导教师签字：________________    日期：____年__月__日")
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(10)

    document.add_paragraph()
    _add_callout(
        document,
        "使用边界",
        "AI 评分仅为辅助建议。本样例中的分数、评语和学情诊断必须经过教师终审后方可用于演示发布，不用于真实教育评价。",
    )


def build_input_report() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    _configure_styles(document)
    _add_header_footer(document)
    document.core_properties.title = "格物智评匿名合成实验报告"
    document.core_properties.subject = "GOAI 2026 AI+教育公开演示材料"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.keywords = "synthetic, anonymized, lab report, demo"
    document.core_properties.comments = "This document contains only synthetic data."

    with tempfile.TemporaryDirectory(prefix="labtrace-doc-") as temp_dir:
        chart_path = Path(temp_dir) / "synthetic_calibration.png"
        _make_chart(chart_path)
        _add_cover(document)
        _add_method_page(document)
        _add_results_page(document, chart_path)
        _add_analysis_page(document)
        document.save(INPUT_PATH)


def build_graded_report() -> dict:
    sys.path.insert(0, str(ROOT))
    from agent_skills.report_injector.scripts.inject_grading_to_docx import inject_all

    source_document = Document(str(INPUT_PATH))
    grading_table_index = next(
        index
        for index, table in enumerate(source_document.tables)
        if "成绩评定"
        in "\n".join(cell.text for row in table.rows for cell in row.cells)
    )
    config = {
        "author": "LabTrace 演示教师",
        "initials": "LT",
        "table_index": grading_table_index,
        "annotations": [
            {
                "text": "目标与原理较清楚；建议补充线性近似适用区间与假设条件。",
                "target": {"type": "keyword", "keyword": "实验目标："},
            },
            {
                "text": "过程能够复现主流程，但应记录恒温槽实际波动范围。",
                "target": {"type": "keyword", "keyword": "实验环境："},
            },
            {
                "text": "表格提供了重复测量与预测值，是支撑得分的关键数据证据。",
                "target": {"type": "keyword", "keyword": "实验数据："},
            },
            {
                "text": "此处只做了定性解释。请计算重复测量标准差或满量程误差，并解释 50 ℃异常点。",
                "target": {"type": "keyword", "keyword": "结果分析："},
            },
            {
                "text": "改进方向具体，但正式结论应直接引用表 1 或图 1 的数据。",
                "target": {"type": "keyword", "keyword": "实验结论："},
            },
        ],
        "scores": [13, 16, 20, 10, 11, 4],
        "comment": (
            "报告完成了主要实验流程，目标、步骤和结果证据较清楚。数据表能够支撑基本结论，"
            "但结果图缺少不确定性表达，误差分析仍偏定性。建议补充重复实验的离散程度、"
            "异常点解释和对照验证，并让结论逐项引用前文数据。本结果为 AI 辅助建议，"
            "正式成绩由教师复核确认。"
        ),
    }
    return inject_all(str(INPUT_PATH), str(GRADED_PATH), config)


def main() -> None:
    build_input_report()
    result = build_graded_report()
    if not all(
        (
            result.get("annotations_count") == 5,
            result.get("score_injected"),
            result.get("comment_injected"),
        )
    ):
        raise RuntimeError(f"graded fixture injection failed: {result}")
    print(f"Created {INPUT_PATH}")
    print(f"Created {GRADED_PATH}")


if __name__ == "__main__":
    main()
