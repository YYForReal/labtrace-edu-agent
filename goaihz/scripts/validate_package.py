#!/usr/bin/env python3
"""Validate the GOAI competition profile without network or model access."""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
REPO_ROOT = ROOT.parent
sys.path.insert(0, str(ROOT / "src"))

from labtrace.contracts import GradeTrace  # noqa: E402
from labtrace.privacy import find_sensitive_data  # noqa: E402


def load_json(path: Path):
    with path.open("r", encoding="utf-8") as stream:
        return json.load(stream)


def validate_project() -> list[str]:
    errors = []
    project = load_json(ROOT / "project.json")
    required = {
        "project_name",
        "track",
        "topic",
        "target_users",
        "closed_loop",
        "source_project",
        "new_contributions",
        "human_decision_boundary",
        "data_policy",
        "license_status",
    }
    missing = sorted(required - set(project))
    if missing:
        errors.append(f"project.json missing fields: {missing}")
    if len(project.get("closed_loop", [])) < 5:
        errors.append("closed_loop must describe a real multi-step workflow")
    if not project.get("source_project", {}).get("base_commit"):
        errors.append("source project base commit must be disclosed")
    if "Apache-2.0" not in project.get("license_status", ""):
        errors.append("standalone public repository must disclose Apache-2.0")
    for name in ("LICENSE", "NOTICE", "SECURITY.md"):
        if not (REPO_ROOT / name).is_file():
            errors.append(f"missing repository policy file: {name}")
    return errors


def validate_rubric() -> list[str]:
    errors = []
    rubric = load_json(ROOT / "config" / "rubrics" / "general_lab_report_v1.json")
    criteria = rubric.get("criteria", [])
    max_total = sum(float(item["max_score"]) for item in criteria)
    weight_total = sum(float(item["weight"]) for item in criteria)
    if abs(max_total - float(rubric.get("total_score", 0))) > 0.01:
        errors.append(
            f"rubric max scores total {max_total}, expected {rubric.get('total_score')}"
        )
    if abs(weight_total - 1) > 0.0001:
        errors.append(f"rubric weights total {weight_total}, expected 1")
    ids = [item["id"] for item in criteria]
    if len(ids) != len(set(ids)):
        errors.append("rubric criterion IDs must be unique")
    return errors


def validate_synthetic_data() -> list[str]:
    errors = []
    data_dir = ROOT / "data" / "synthetic"
    for path in sorted(data_dir.glob("*.json")):
        text = path.read_text(encoding="utf-8")
        findings = find_sensitive_data(text)
        if findings:
            errors.append(f"{path.name} contains possible personal data: {findings}")
    GradeTrace.from_dict(load_json(data_dir / "trace_case.json"))
    for path in sorted(data_dir.glob("*.docx")):
        document = Document(str(path))
        text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ]
        )
        findings = find_sensitive_data(text)
        if findings:
            errors.append(f"{path.name} contains possible personal data: {findings}")
        properties = document.core_properties
        metadata = {
            "author": properties.author,
            "last_modified_by": properties.last_modified_by,
        }
        allowed_metadata = {"LabTrace Open Source Demo"}
        populated = {
            key: value
            for key, value in metadata.items()
            if value and value not in allowed_metadata
        }
        if populated:
            errors.append(
                f"{path.name} contains identifying core metadata: {populated}"
            )
    return errors


def validate_required_docs() -> list[str]:
    required = {
        "product.md",
        "closed_loop_demo_plan.md",
        "engineering_plan_v1.md",
        "architecture.md",
        "compliance.md",
        "evaluation.md",
        "preliminary_submission.md",
        "demo_script.md",
        "competition_notes.md",
        "team_names.md",
        "open_source_boundary.md",
    }
    existing = {path.name for path in (ROOT / "docs").glob("*.md")}
    missing = sorted(required - existing)
    public_assets = {
        "demo-assignment-template_实验任务书.docx",
        "demo-allergen-001_实验报告.docx",
        "demo-game-dev-001_实验报告.docx",
    }
    missing.extend(
        sorted(
            f"data/synthetic/{name}"
            for name in public_assets
            if not (ROOT / "data" / "synthetic" / name).is_file()
        )
    )
    return [f"missing required docs: {missing}"] if missing else []


def validate_submission() -> list[str]:
    errors = []
    submission_dir = ROOT / "submission"
    required = {
        "作品简介.txt",
        "格物智评_LabTrace_GOAI初赛方案.pptx",
        "格物智评_LabTrace_GOAI初赛方案.pdf",
        "格物智评_LabTrace_初赛Demo.mp4",
        "格物智评_LabTrace_Demo.zh-CN.srt",
        "初赛作品说明.md",
        "初赛提交清单.md",
        "演示运行说明.md",
        "数据与合规说明.md",
    }
    missing = sorted(
        name
        for name in required
        if not (submission_dir / name).is_file()
        or (submission_dir / name).stat().st_size == 0
    )
    if missing:
        errors.append(f"missing or empty submission files: {missing}")
    intro_path = submission_dir / "作品简介.txt"
    if intro_path.exists():
        intro = intro_path.read_text(encoding="utf-8").strip()
        if len(intro) > 500:
            errors.append(f"作品简介 is {len(intro)} characters; expected <= 500")
    return errors


def main() -> int:
    checks = {
        "project": validate_project,
        "rubric": validate_rubric,
        "synthetic_data": validate_synthetic_data,
        "documents": validate_required_docs,
        "submission": validate_submission,
    }
    all_errors = []
    for name, check in checks.items():
        try:
            errors = check()
        except Exception as exc:
            errors = [f"{type(exc).__name__}: {exc}"]
        if errors:
            all_errors.extend(f"{name}: {error}" for error in errors)
            print(f"[FAIL] {name}")
        else:
            print(f"[ OK ] {name}")

    if all_errors:
        print("\nValidation errors:")
        for error in all_errors:
            print(f"- {error}")
        return 1

    print("\nGOAI competition profile is internally consistent.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
