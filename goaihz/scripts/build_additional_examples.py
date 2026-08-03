"""Build two public, identity-free LabTrace examples.

The allergen report is fully synthetic. The game-development report is a
synthetic reconstruction based on recurring structures observed in a real
course corpus; it does not reuse student prose, screenshots, names, IDs, or
document metadata.
"""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from build_synthetic_documents import (
    DOC_FONT,
    MUTED,
    NAVY,
    ORANGE,
    TEAL,
    _add_callout,
    _add_header_footer,
    _add_section_title,
    _configure_styles,
    _set_cell_margins,
    _set_cell_shading,
    _set_repeat_table_header,
    _set_table_width,
)


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "goaihz" / "data" / "synthetic"
ALLERGEN_PATH = OUTPUT_DIR / "demo-allergen-001_实验报告.docx"
GAME_PATH = OUTPUT_DIR / "demo-game-dev-001_实验报告.docx"
FONT_PATH = ROOT / "goaihz" / "assets" / "fonts" / "FandolHei-Regular.otf"


def _font(size: int):
    return ImageFont.truetype(str(FONT_PATH), size)


def _new_document(title: str, subject: str) -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    _configure_styles(document)
    _add_header_footer(document)
    properties = document.core_properties
    properties.title = title
    properties.subject = subject
    properties.author = ""
    properties.last_modified_by = ""
    properties.keywords = "synthetic, anonymized, labtrace, public demo"
    properties.comments = "Identity-free public demonstration fixture."
    return document


def _cover(
    document: Document,
    *,
    domain: str,
    title: str,
    subtitle: str,
    case_id: str,
    disclosure: str,
) -> None:
    kicker = document.add_paragraph()
    kicker.add_run("PUBLIC DEMO / 身份无关演示材料").font.color.rgb = (
        RGBColor.from_string(ORANGE)
    )
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(48)
    run = heading.add_run(domain)
    run.font.name = DOC_FONT
    run.font.size = Pt(17)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)
    title_paragraph = document.add_paragraph(style="Title")
    title_paragraph.add_run(title)
    subtitle_paragraph = document.add_paragraph(subtitle)
    subtitle_paragraph.paragraph_format.space_after = Pt(28)
    for run in subtitle_paragraph.runs:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    metadata = document.add_table(rows=5, cols=2)
    _set_table_width(metadata)
    rows = (
        ("课程", f"{domain}实验方法（公开演示课程）"),
        ("用例编号", case_id),
        ("提交身份", "匿名学生样例（无真实身份）"),
        ("报告日期", "2026 年 7 月 26 日"),
        ("使用范围", "GOAI 2026 · AI+教育 Demo"),
    )
    for index, (label, value) in enumerate(rows):
        left, right = metadata.rows[index].cells
        _set_cell_shading(left, "EEF3F2")
        left.paragraphs[0].add_run(label).bold = True
        right.paragraphs[0].add_run(value)
    document.add_paragraph()
    _add_callout(document, "来源声明", disclosure)
    document.add_page_break()


def _data_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    _set_table_width(table)
    for cell, value in zip(table.rows[0].cells, headers):
        _set_cell_shading(cell, NAVY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    _set_repeat_table_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            cell.text = value
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_margins(cell, top=70, bottom=70)


def _add_review_area(document: Document) -> None:
    _add_section_title(document, "06", "教师批阅区")
    table = document.add_table(rows=1, cols=1)
    _set_table_width(table)
    cell = table.cell(0, 0)
    _set_cell_shading(cell, "F7F9F8")
    cell.paragraphs[0].add_run("指导教师批阅意见：").bold = True
    cell.add_paragraph("成绩评定：")
    cell.add_paragraph("评语：")
    cell.add_paragraph("指导教师签字：________________    日期：____年__月__日")
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(10)
    document.add_paragraph()
    _add_callout(
        document,
        "教育边界",
        "本用例只验证证据化批改流程。AI 输出是辅助建议，正式成绩和反馈必须由教师确认。",
    )


def _allergen_chart(path: Path) -> None:
    image = Image.new("RGB", (1280, 650), "#F7F9F8")
    draw = ImageDraw.Draw(image)
    draw.text(
        (68, 50), "Ara h 1 标准曲线（完全合成教学数据）", font=_font(36), fill="#173E59"
    )
    left, top, right, bottom = 115, 135, 1190, 540
    for index in range(6):
        y = bottom - index * (bottom - top) / 5
        draw.line((left, y, right, y), fill="#D9E4E1", width=2)
        draw.text((58, y - 11), f"{index * 0.5:.1f}", font=_font(18), fill="#62716B")
    concentrations = [0, 2, 5, 10, 20, 40]
    absorbance = [0.05, 0.19, 0.43, 0.81, 1.51, 2.42]
    points = []
    for concentration, value in zip(concentrations, absorbance):
        x = left + concentration / 40 * (right - left)
        y = bottom - value / 2.5 * (bottom - top)
        points.append((x, y))
    draw.line(points, fill="#168A82", width=6)
    for x, y in points:
        draw.ellipse(
            (x - 9, y - 9, x + 9, y + 9), fill="#D66443", outline="white", width=3
        )
    draw.text((505, 575), "标准品浓度 / ng/mL", font=_font(22), fill="#4A5B54")
    draw.text((38, 112), "OD450", font=_font(20), fill="#4A5B54")
    image.save(path)


def build_allergen_report() -> None:
    document = _new_document(
        "过敏原蛋白 ELISA 检测实验报告",
        "完全合成生命科学教学案例，不涉及患者和医学诊断",
    )
    _cover(
        document,
        domain="生命科学实验",
        title="过敏原蛋白 ELISA\n定量检测",
        subtitle="用标准曲线验证证据链与教师复核的完全合成案例",
        case_id="DEMO-ALLERGEN-001",
        disclosure=(
            "本报告的样本、数值、图表和身份均为人工构造。案例只模拟花生过敏原蛋白"
            " Ara h 1 的教学检测，不对应患者、食品安全结论或医学诊断。"
        ),
    )
    _add_section_title(document, "01", "实验目标与原理")
    document.add_paragraph(
        "实验目标：使用夹心 ELISA 的合成教学数据建立 Ara h 1 标准曲线，估算三个模拟食品样本的蛋白浓度，并检查空白、阳性和基质对照是否支持结果解释。"
    )
    document.add_paragraph(
        "实验原理：样本中的目标蛋白与包被抗体结合，再由检测抗体和显色反应形成 OD450 信号。在限定区间内，信号随浓度增加而上升，可通过标准曲线完成教学用定量估算。"
    )
    _add_section_title(document, "02", "实验方法与过程")
    document.add_paragraph(
        "实验环境：96 孔板、教学用模拟标准品、空白对照、阳性对照和三种模拟食品提取液；每个标准点设置两个复孔，显色 12 min 后读取 OD450。"
    )
    document.add_paragraph(
        "实验方法：依次完成加样、孵育、洗板、检测抗体、显色和终止步骤。标准品浓度为 0、2、5、10、20、40 ng·mL⁻¹；样本结果按同一批次标准曲线换算。"
    )
    _data_table(
        document,
        ["阶段", "关键控制", "记录", "证据"],
        [
            ["标准品", "六点、双复孔", "浓度与 OD450", "标准曲线"],
            ["空白", "不加目标蛋白", "OD450 < 0.10", "背景检查"],
            ["阳性", "已知教学浓度", "回收率", "流程有效性"],
            ["样本", "统一稀释倍数", "双复孔均值", "浓度估算"],
        ],
    )
    document.add_page_break()
    _add_section_title(document, "03", "实验数据与结果")
    _data_table(
        document,
        ["对象", "复孔 1", "复孔 2", "均值", "估算浓度"],
        [
            ["空白", "0.05", "0.06", "0.055", "未检出"],
            ["阳性对照", "0.82", "0.79", "0.805", "约 10.0 ng·mL⁻¹"],
            ["模拟样本 A", "0.11", "0.13", "0.120", "约 1.0 ng·mL⁻¹"],
            ["模拟样本 B", "1.18", "1.23", "1.205", "约 15.5 ng·mL⁻¹"],
            ["模拟样本 C", "2.31", "2.38", "2.345", "约 38.5 ng·mL⁻¹"],
        ],
    )
    document.add_paragraph("表 1  模拟样本 OD450 与浓度估算（完全合成）").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    with tempfile.TemporaryDirectory(prefix="labtrace-allergen-") as temp_dir:
        chart_path = Path(temp_dir) / "allergen_curve.png"
        _allergen_chart(chart_path)
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(chart_path), width=Inches(6.35))
        document.add_paragraph("图 1  合成标准曲线").alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )
    document.add_paragraph(
        "实验数据：标准曲线随浓度单调上升，空白低于 0.10，阳性对照接近教学设定值。模拟样本 B、C 的信号明显高于样本 A。"
    )
    document.add_page_break()
    _add_section_title(document, "04", "结果分析与验证")
    document.add_paragraph(
        "结果分析：空白和阳性对照表现正常，说明本次模拟流程基本有效。样本 C 的估算浓度最高，样本 A 接近曲线低端。"
    )
    document.add_paragraph(
        "误差分析：当前报告只比较了两个复孔，没有计算变异系数，也没有进行加标回收和不同稀释倍数验证。样本 C 接近曲线上限，是否需要重新稀释仍应由教师结合课程要求判断。"
    )
    _add_section_title(document, "05", "实验结论与反思")
    document.add_paragraph(
        "实验结论：本次教学实验完成了标准曲线和三个模拟样本的浓度估算，对照结果支持基本流程有效。"
    )
    document.add_paragraph(
        "改进方向：增加复孔数量并计算变异系数；对高值样本追加稀释复测；用加标回收验证基质影响。任何结果都不得用于真实过敏诊断或食品安全判断。"
    )
    _add_review_area(document)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(ALLERGEN_PATH)


def _game_diagram(path: Path) -> None:
    image = Image.new("RGB", (1280, 650), "#16231F")
    draw = ImageDraw.Draw(image)
    draw.text(
        (58, 44), "Unity 弹射原型 · 运行证据重构示意", font=_font(34), fill="#F4F1E8"
    )
    draw.rounded_rectangle((65, 130, 350, 540), radius=22, fill="#214D41")
    draw.rounded_rectangle((915, 175, 1205, 510), radius=18, fill="#D9C7A4")
    draw.ellipse((185, 320, 235, 370), fill="#E46C47")
    draw.line((230, 344, 900, 290), fill="#F4A384", width=5)
    for offset in (0, 72, 144):
        draw.rectangle(
            (945 + offset, 390 - offset // 3, 1005 + offset, 510), fill="#526D63"
        )
    draw.rectangle((1080, 250, 1140, 315), fill="#E46C47")
    draw.text((95, 565), "发射点 / 连续碰撞检测", font=_font(20), fill="#B9CBC4")
    draw.text((900, 565), "目标区 / 命中事件日志", font=_font(20), fill="#B9CBC4")
    image.save(path)


def build_game_report() -> None:
    document = _new_document(
        "Unity 弹射原型与碰撞验证实验报告",
        "基于实际课程结构合成重构的游戏开发教学案例",
    )
    _cover(
        document,
        domain="游戏开发实验",
        title="Unity 弹射原型\n与碰撞验证",
        subtitle="来自实际课程任务结构的合成重构案例",
        case_id="DEMO-GAME-DEV-001",
        disclosure=(
            "本公开报告参考实际游戏开发课程常见任务结构与证据类型重新编写。"
            "不复用学生姓名、学号、原文、截图、代码仓库、教师信息或文档元数据。"
        ),
    )
    _add_section_title(document, "01", "实验目标与原理")
    document.add_paragraph(
        "实验目标：在 Unity 中完成一个可重复测试的弹射原型，验证发射速度、连续碰撞检测、目标命中事件和摄像机跟随是否形成完整运行闭环。"
    )
    document.add_paragraph(
        "实验原理：弹丸由 Rigidbody 驱动物理运动，碰撞器负责接触判断，目标触发器记录命中事件；摄像机使用插值跟随兴趣点，以避免视角瞬间跳变。"
    )
    _add_section_title(document, "02", "实验方法与过程")
    document.add_paragraph(
        "实验环境：Unity 2022 LTS、C#、固定 16:9 游戏视图；弹丸质量 5，Collision Detection 设为 Continuous，冻结 Z 轴位移。"
    )
    document.add_paragraph(
        "实验方法：先搭建发射点、弹丸和目标，再实现鼠标瞄准与发射；随后加入目标命中事件、轨迹渲染和摄像机平滑跟随，最后用三组速度参数完成重复运行测试。"
    )
    _data_table(
        document,
        ["模块", "实现", "验证动作", "输出证据"],
        [
            ["弹射", "Rigidbody + 发射向量", "三组速度运行", "落点记录"],
            ["碰撞", "Continuous + Trigger", "高速命中", "事件日志"],
            ["摄像机", "Vector3.Lerp", "全程跟随", "视图截图"],
            ["关卡", "目标与重置", "命中后切换", "状态文本"],
        ],
    )
    document.add_page_break()
    _add_section_title(document, "03", "实验数据与结果")
    _data_table(
        document,
        ["测试", "初速度", "是否命中", "事件日志", "观察"],
        [
            ["T1", "18", "否", "未触发", "落点偏近"],
            ["T2", "22", "是", "GoalHit=1", "轨迹稳定"],
            ["T3", "26", "是", "GoalHit=1", "摄像机缩放偏快"],
        ],
    )
    document.add_paragraph("表 1  弹射原型三次运行记录（合成重构）").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    with tempfile.TemporaryDirectory(prefix="labtrace-game-") as temp_dir:
        diagram_path = Path(temp_dir) / "game_evidence.png"
        _game_diagram(diagram_path)
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(diagram_path), width=Inches(6.35))
        document.add_paragraph("图 1  运行证据重构示意（非学生原截图）").alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )
    document.add_paragraph(
        "实验数据：T2 与 T3 均触发目标事件，说明碰撞和关卡状态链路已经接通；T1 未命中，为调整发射参数提供了反例。"
    )
    document.add_page_break()
    _add_section_title(document, "04", "结果分析与验证")
    document.add_paragraph(
        "结果分析：连续碰撞检测减少了高速弹丸穿透目标的情况，22 与 26 两组速度都能完成命中。摄像机能够跟随弹丸，但高速度下缩放变化较快。"
    )
    document.add_paragraph(
        "误差分析：当前只有三次运行，尚未覆盖不同帧率、极端速度、目标边缘碰撞和连续重置。报告也没有统计命中率或记录帧时间，因此“运行稳定”的结论仍偏定性。"
    )
    _add_section_title(document, "05", "实验结论与反思")
    document.add_paragraph(
        "实验结论：弹射、碰撞、目标事件和摄像机跟随已经形成最小可玩闭环，主要功能能够运行。"
    )
    document.add_paragraph(
        "改进方向：增加十次以上重复测试，记录命中率和帧时间；补充边缘碰撞、极端速度和重置压力测试，并让结论逐项引用日志与数据。"
    )
    _add_review_area(document)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(GAME_PATH)


def main() -> None:
    build_allergen_report()
    build_game_report()
    print(f"Created {ALLERGEN_PATH}")
    print(f"Created {GAME_PATH}")


if __name__ == "__main__":
    main()
