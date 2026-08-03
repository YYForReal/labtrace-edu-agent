#!/usr/bin/env python3
"""Build the synthetic public assignment brief distributed with the demo."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "data" / "synthetic" / "demo-assignment-template_实验任务书.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "24364B"
MUTED = "66737F"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF4D6"
WHITE = "FFFFFF"


def set_run_font(
    run,
    *,
    size: float = 11,
    bold: bool = False,
    color: str = INK,
    italic: bool = False,
):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_fill(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = 120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_table(
    doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_fill(cell, PALE_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        set_run_font(paragraph.add_run(value), bold=True, color=DARK_BLUE)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            if index == len(values) - 1 and len(value) <= 8:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(paragraph.add_run(value), size=10.2)
        set_table_geometry(table, widths)
    return table


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(
        paragraph.add_run("LabTrace 公开合成模板  |  第 "), size=8.5, color=MUTED
    )
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    set_run_font(paragraph.add_run(" 页"), size=8.5, color=MUTED)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Title": (28, INK, 0, 8),
        "Subtitle": (14, MUTED, 0, 22),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name not in {"Subtitle"}
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Number", "List Bullet"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build() -> Path:
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    set_run_font(header.add_run("高校实验课程 · 公开演示任务书"), size=8.5, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(8)
    set_run_font(
        kicker.add_run("COMPUTER GAME DEVELOPMENT LAB"), size=9, bold=True, color=BLUE
    )
    doc.add_heading("计算机游戏开发实验任务书", level=0)
    subtitle = doc.add_paragraph(style="Subtitle")
    set_run_font(
        subtitle.add_run("实验一：Unity 弹射原型与碰撞验证"), size=14, color=MUTED
    )

    badge = doc.add_paragraph()
    badge.paragraph_format.space_after = Pt(18)
    set_run_font(
        badge.add_run("人工合成公开模板 · 不含真实课程、教师或学生信息"),
        size=9.5,
        bold=True,
        color="7A5A00",
    )

    add_table(
        doc,
        ["建议课时", "实验环境", "核心交付", "评价方式"],
        [
            [
                "4 学时",
                "Unity 2022.3 LTS / C#",
                "图文实验报告 DOCX",
                "AI 建议 + 教师终审",
            ]
        ],
        [1440, 2520, 2880, 2520],
    )

    doc.add_heading("一、任务背景", level=1)
    doc.add_paragraph(
        "实现一个可运行的弹射原型：玩家通过输入控制发射方向与力度，刚体在重力作用下运动，"
        "与目标或地面发生碰撞后给出反馈，并能复位进入下一轮测试。实验重点不是展示最终游戏美术，"
        "而是用可复核证据说明物理参数、碰撞逻辑和验证过程。"
    )

    doc.add_heading("二、学习目标", level=1)
    goals = [
        "解释 Rigidbody、Collider、重力、冲量与碰撞回调在弹射系统中的作用。",
        "用脚本实现输入、发射、碰撞反馈和复位的完整交互闭环。",
        "记录关键参数与至少三次运行结果，并用截图或表格支撑结论。",
        "识别高速穿透、镜头抖动或复位异常等边界问题，提出可验证的改进方案。",
    ]
    for item in goals:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("三、实验任务", level=1)
    tasks = [
        (
            "1. 场景与组件",
            "建立发射体、目标、地面与摄像机；记录对象层级、Collider 类型和 Rigidbody 关键参数。",
        ),
        (
            "2. 弹射控制",
            "实现方向与力度输入，以冲量或等价方式发射；说明输入值如何映射到物理量。",
        ),
        (
            "3. 碰撞与复位",
            "使用碰撞回调记录命中/落地结果，避免重复计分，并让系统能够稳定复位。",
        ),
        (
            "4. 运行验证",
            "至少执行三组参数测试，包含一组边界条件；保留运行截图、数据表和异常分析。",
        ),
    ]
    for title, body in tasks:
        doc.add_heading(title, level=2)
        doc.add_paragraph(body)

    doc.add_heading("四、报告证据清单", level=1)
    add_table(
        doc,
        ["证据项", "最低要求", "建议放置位置"],
        [
            [
                "环境与参数",
                "Unity 版本、主要组件、质量/阻力/碰撞检测等关键参数",
                "方法与过程",
            ],
            [
                "核心代码",
                "发射、碰撞或复位的关键片段，并解释输入和状态变化",
                "方法与过程",
            ],
            [
                "运行截图",
                "至少 1 张能识别发射体、目标和运行状态的合成或授权图片",
                "数据与结果",
            ],
            ["测试记录", "至少 3 次运行，记录力度、结果、异常和观察", "数据与结果"],
            [
                "分析结论",
                "区分“运行成功”与“验证充分”，讨论至少 1 个边界问题",
                "分析与反思",
            ],
        ],
        [2160, 5040, 2160],
    )

    doc.add_heading("五、评分参考（100 分）", level=1)
    add_table(
        doc,
        ["维度", "评价重点", "分值"],
        [
            ["实验目标与原理", "目标、物理原理与组件关系", "15"],
            ["实验方法与过程", "环境、参数、步骤和可复现性", "20"],
            ["数据、结果与证据", "运行数据、截图、代码和结果对应", "25"],
            ["分析、验证与误差讨论", "边界测试、异常解释和验证充分性", "20"],
            ["结论与反思", "回应目标并提出可验证的改进", "15"],
            ["报告规范与学术表达", "结构、图表、单位、引用和来源边界", "5"],
        ],
        [2520, 5760, 1080],
    )

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(10)
    note.paragraph_format.left_indent = Inches(0.12)
    set_run_font(
        note.add_run(
            "教师终审边界：AI 评分仅为建议。证据不足、低置信度或图片判断必须由教师核对；"
            "最终成绩和对学生发布的评语由教师确认。"
        ),
        size=10.5,
        bold=True,
        color="7A5A00",
    )
    p_pr = note._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE_GOLD)
    p_pr.append(shd)

    doc.add_heading("六、提交前自检", level=1)
    checks = [
        "报告中的每个主要结论都能引用到正文、表格、代码或图片证据。",
        "图片来源和处理授权已经明确，不包含姓名、学号、头像或其他身份信息。",
        "测试记录包含成功与失败/边界情况，而不是只展示一次成功截图。",
        "DOCX 可以正常打开，图表清晰，章节和单位完整。",
    ]
    for item in checks:
        doc.add_paragraph(item, style="List Bullet")

    props = doc.core_properties
    props.title = "Unity 弹射原型与碰撞验证实验任务书"
    props.subject = "LabTrace 公开合成演示模板"
    props.author = "LabTrace Open Source Demo"
    props.last_modified_by = "LabTrace Open Source Demo"
    props.keywords = "synthetic, education, lab report, Unity"
    props.comments = "人工合成公开模板，不含真实课程或个人信息。"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
