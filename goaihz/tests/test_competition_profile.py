from __future__ import annotations

import json
import io
import os
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

ROOT = Path(__file__).resolve().parents[1]
REPO_ROOT = ROOT.parent
sys.path.insert(0, str(ROOT / "src"))
sys.path.insert(0, str(REPO_ROOT))

from app.config import AppConfig
from app.agent.llm_client import LLMResponse, TextBlock
from app.config import LLMConfig
from labtrace.contracts import ContractError, GradeTrace, ReviewDecision
from labtrace.diagnosis import build_class_diagnosis
from labtrace.privacy import find_sensitive_data, pseudonymize
from labtrace.rubric import RubricError, validate_rubric
from goaihz import api as demo_api
from goaihz.demo_engine import build_demo_trace, parse_report
from goaihz.model_engine import build_model_evidence, build_model_trace


class CompetitionProfileTests(unittest.TestCase):
    def setUp(self):
        self._previous_llm_enabled = os.environ.get("LABTRACE_LLM_ENABLED")
        os.environ["LABTRACE_LLM_ENABLED"] = "false"
        demo_api._tasks.clear()
        demo_api._grade_attempts.clear()

    def tearDown(self):
        if self._previous_llm_enabled is None:
            os.environ.pop("LABTRACE_LLM_ENABLED", None)
        else:
            os.environ["LABTRACE_LLM_ENABLED"] = self._previous_llm_enabled

    def load_json(self, relative_path: str):
        return json.loads((ROOT / relative_path).read_text(encoding="utf-8"))

    def test_competition_rubric_totals_100(self):
        rubric = self.load_json("config/rubrics/general_lab_report_v1.json")
        self.assertEqual(sum(item["max_score"] for item in rubric["criteria"]), 100)
        self.assertAlmostEqual(sum(item["weight"] for item in rubric["criteria"]), 1)
        normalized = validate_rubric(rubric)
        self.assertEqual(normalized["total_score"], 100)

    def test_teacher_rubric_rejects_inconsistent_total(self):
        rubric = self.load_json("config/rubrics/general_lab_report_v1.json")
        rubric["total_score"] = 99
        with self.assertRaises(RubricError):
            validate_rubric(rubric)

    def test_app_config_accepts_competition_rubric_directory(self):
        expected = str(ROOT / "config" / "rubrics")
        with patch.dict(os.environ, {"RUBRICS_DIR": expected}):
            self.assertEqual(AppConfig().rubrics_dir, expected)

    def test_sample_trace_is_valid_and_auditable(self):
        trace = GradeTrace.from_dict(self.load_json("data/synthetic/trace_case.json"))
        self.assertTrue(trace.needs_human_review)
        self.assertEqual(trace.model_total_score, 33)

    def test_unknown_evidence_reference_is_rejected(self):
        data = self.load_json("data/synthetic/trace_case.json")
        data["criteria"][0]["evidence_ids"] = ["missing"]
        with self.assertRaises(ContractError):
            GradeTrace.from_dict(data)

    def test_diagnosis_finds_analysis_as_top_weakness(self):
        records = self.load_json("data/synthetic/grade_records.json")
        result = build_class_diagnosis(records)
        self.assertEqual(result["record_count"], 3)
        self.assertEqual(
            result["top_weaknesses"][0]["criterion_id"],
            "analysis_and_validation",
        )

    def test_privacy_scanner_and_pseudonymizer(self):
        text = "张同学 13800138000 zhang@example.edu 11010519491231002X"
        kinds = {item["kind"] for item in find_sensitive_data(text)}
        self.assertEqual(kinds, {"email", "phone_cn", "national_id_cn"})
        redacted = pseudonymize(text, student_name="张同学")
        self.assertNotIn("张同学", redacted)
        self.assertNotIn("13800138000", redacted)

    def test_model_evidence_redacts_identity_before_external_call(self):
        parsed = {
            "file_path": "/private/张三-20260001.docx",
            "student_info": {"name": "张三", "student_id": "20260001"},
            "paragraphs": [
                {
                    "text": (
                        "张三 20260001 联系方式 13800138000，"
                        "邮箱 zhangsan@example.edu，完成了实验。"
                    )
                }
            ],
            "tables": [],
            "images": [],
        }
        evidence, privacy = build_model_evidence(parsed)
        serialized = json.dumps([item.excerpt for item in evidence], ensure_ascii=False)
        self.assertNotIn("张三", serialized)
        self.assertNotIn("20260001", serialized)
        self.assertNotIn("13800138000", serialized)
        self.assertNotIn("zhangsan@example.edu", serialized)
        self.assertGreaterEqual(privacy["detected_sensitive_items"], 4)

    def test_real_model_output_maps_to_same_validated_trace_contract(self):
        rubric = validate_rubric(
            self.load_json("config/rubrics/general_lab_report_v1.json")
        )
        parsed = parse_report(
            str(ROOT / "data" / "synthetic" / "demo-game-dev-001_实验报告.docx")
        )
        criteria = []
        for item in rubric["criteria"]:
            criteria.append(
                {
                    "criterion_id": item["id"],
                    "score": item["max_score"] * 0.7,
                    "reason": "该维度有可定位证据，但仍需教师核对完成质量。",
                    "evidence_ids": ["p-0001"],
                    "confidence": 0.8,
                }
            )
        raw = json.dumps(
            {
                "criteria": criteria,
                "overall_summary": "完成主要实验流程。",
                "risks": ["图片仅使用邻近文本，未直接进行视觉判断。"],
            },
            ensure_ascii=False,
        )

        class FakeClient:
            def create_message(self, **_kwargs):
                return LLMResponse(
                    text_blocks=[TextBlock(raw)],
                    input_tokens=100,
                    output_tokens=200,
                )

        config = LLMConfig(
            api_key="test-only",
            model="MiniMax-M3",
            enable_thinking=False,
            failover_endpoints=[],
        )
        trace, run, privacy = build_model_trace(
            parsed,
            trace_id="trace-fake-model",
            submission_alias="submission-fake-model",
            rubric=rubric,
            client=FakeClient(),
            config=config,
        )
        self.assertEqual(len(trace.criteria), len(rubric["criteria"]))
        self.assertTrue(trace.needs_human_review)
        self.assertTrue(run["structured_output_validated"])
        self.assertEqual(run["tokens"], {"input": 100, "output": 200})
        self.assertFalse(privacy["images_sent_to_text_model"])

    def test_opt_in_image_maps_parser_index_to_existing_evidence_id(self):
        rubric = validate_rubric(
            self.load_json("config/rubrics/general_lab_report_v1.json")
        )
        parsed = parse_report(
            str(ROOT / "data" / "synthetic" / "demo-game-dev-001_实验报告.docx")
        )
        self.assertEqual(parsed["images_for_vision"][0]["index"], 0)
        raw = json.dumps(
            {
                "criteria": [
                    {
                        "criterion_id": item["id"],
                        "score": item["max_score"] * 0.6,
                        "reason": "已依据文本和授权图片提出建议，仍需教师终审。",
                        "evidence_ids": ["p-0001", "i-0001"],
                        "confidence": 0.75,
                    }
                    for item in rubric["criteria"]
                ],
                "risks": ["图片只在本次任务授权范围内处理。"],
            },
            ensure_ascii=False,
        )

        class CapturingClient:
            request = None

            def create_message(self, **kwargs):
                self.request = kwargs
                return LLMResponse(
                    text_blocks=[TextBlock(raw)],
                    input_tokens=110,
                    output_tokens=210,
                )

        client = CapturingClient()
        trace, run, privacy = build_model_trace(
            parsed,
            trace_id="trace-image-index",
            submission_alias="submission-image-index",
            rubric=rubric,
            allow_external_images=True,
            client=client,
            config=LLMConfig(
                api_key="test-only",
                model="MiniMax-M3",
                enable_thinking=False,
                failover_endpoints=[],
            ),
        )
        serialized_request = json.dumps(client.request, ensure_ascii=False)
        self.assertIn("evidence_id=i-0001", serialized_request)
        self.assertNotIn("evidence_id=i-0000", serialized_request)
        self.assertEqual(run["images_sent"], 1)
        self.assertEqual(privacy["external_image_count"], 1)
        self.assertEqual(
            next(
                item for item in trace.evidence if item.evidence_id == "i-0001"
            ).verification,
            "model_observed",
        )
        image_evidence = next(
            item for item in trace.evidence if item.evidence_id == "i-0001"
        )
        self.assertRegex(image_evidence.locator, r"^image:1@paragraph:\d+$")

    def test_parser_keeps_original_word_anchor_for_embedded_image(self):
        parsed = parse_report(
            str(ROOT / "data" / "synthetic" / "demo-game-dev-001_实验报告.docx")
        )
        image = parsed["images"][0]
        vision_image = parsed["images_for_vision"][0]
        self.assertIsInstance(image["docx_paragraph_index"], int)
        self.assertGreaterEqual(image["docx_paragraph_index"], 0)
        self.assertEqual(
            vision_image["docx_paragraph_index"],
            image["docx_paragraph_index"],
        )

    def test_generic_image_rich_word_gets_native_comments_and_teacher_feedback(self):
        rubric = validate_rubric(
            self.load_json("config/rubrics/general_lab_report_v1.json")
        )
        source_sample = ROOT / "data" / "synthetic" / "demo-game-dev-001_实验报告.docx"
        with tempfile.TemporaryDirectory(prefix="labtrace-word-native-") as temp_dir:
            temp_root = Path(temp_dir)
            generic_input = temp_root / "anonymous-image-report.docx"
            output_path = temp_root / "anonymous-image-report-reviewed.docx"
            with zipfile.ZipFile(source_sample) as package:
                media_name = next(
                    name
                    for name in package.namelist()
                    if name.startswith("word/media/")
                )
                image_blob = package.read(media_name)

            document = Document()
            document.add_heading("游戏开发实验报告", level=1)
            document.add_paragraph("实验目标：验证弹射原型的碰撞、状态切换与边界行为。")
            document.add_paragraph("运行结果：以下图片记录匿名合成的运行界面。")
            document.add_picture(io.BytesIO(image_blob))
            document.add_paragraph(
                "结果分析：原型能够运行，但仍需补充边界条件与批量测试。"
            )
            document.save(generic_input)

            parsed = parse_report(str(generic_input))
            self.assertEqual(parsed["document_profile"]["image_count"], 1)
            raw = json.dumps(
                {
                    "criteria": [
                        {
                            "criterion_id": item["id"],
                            "score": item["max_score"] * 0.7,
                            "reason": "已结合授权图片与正文定位证据，仍需教师终审。",
                            "evidence_ids": ["i-0001"],
                            "confidence": 0.8,
                        }
                        for item in rubric["criteria"]
                    ],
                    "risks": ["图片结论仍需教师结合课程要求确认。"],
                },
                ensure_ascii=False,
            )

            class ImageClient:
                def create_message(self, **_kwargs):
                    return LLMResponse(
                        text_blocks=[TextBlock(raw)],
                        input_tokens=120,
                        output_tokens=220,
                    )

            suggested, _, _ = build_model_trace(
                parsed,
                trace_id="trace-word-native",
                submission_alias="submission-word-native",
                rubric=rubric,
                allow_external_images=True,
                client=ImageClient(),
                config=LLMConfig(
                    api_key="test-only",
                    model="MiniMax-M3",
                    enable_thinking=False,
                    failover_endpoints=[],
                ),
            )
            reviewed = GradeTrace(
                trace_id=suggested.trace_id,
                rubric_id=suggested.rubric_id,
                submission_alias=suggested.submission_alias,
                evidence=suggested.evidence,
                criteria=suggested.criteria,
                model_total_score=suggested.model_total_score,
                needs_human_review=False,
                review_reasons=suggested.review_reasons,
                review=ReviewDecision(
                    status="approved",
                    reviewer_role="teacher",
                    final_score=suggested.model_total_score,
                    note="教师已核对图片、正文和课程标准，同意发布。",
                ),
            )
            reviewed.validate()
            delivery = demo_api._try_build_annotated_report(
                generic_input,
                output_path,
                reviewed,
            )
            self.assertTrue(delivery["available"])
            details = delivery["details"]
            self.assertEqual(details["delivery_mode"], "generic_appendix")
            self.assertEqual(details["annotations_count"], 6)
            self.assertEqual(details["image_annotations_count"], 6)
            self.assertTrue(details["score_injected"])
            self.assertTrue(details["comment_injected"])
            self.assertTrue(details["generic_section_appended"])
            self.assertTrue(details["evidence_appendix_appended"])

            reviewed_document = Document(output_path)
            visible_text = "\n".join(
                [paragraph.text for paragraph in reviewed_document.paragraphs]
                + [
                    cell.text
                    for table in reviewed_document.tables
                    for row in table.rows
                    for cell in row.cells
                ]
            )
            self.assertIn("教师批改意见", visible_text)
            self.assertIn("教师评语", visible_text)
            self.assertIn("教师已核对图片、正文和课程标准", visible_text)
            self.assertIn("附录：证据引用索引（LabTrace）", visible_text)
            self.assertIn("[1]", visible_text)
            with zipfile.ZipFile(output_path) as package:
                self.assertIn("word/comments.xml", package.namelist())
                comments = package.read("word/comments.xml").decode("utf-8")
                document_xml = package.read("word/document.xml").decode("utf-8")
                self.assertEqual(comments.count("<w:comment "), 6)
                self.assertEqual(document_xml.count("<w:commentRangeStart"), 6)
                self.assertTrue(
                    any(name.startswith("word/media/") for name in package.namelist())
                )

    def test_model_contract_error_is_explained_and_repaired_once(self):
        rubric = validate_rubric(
            self.load_json("config/rubrics/general_lab_report_v1.json")
        )
        parsed = parse_report(
            str(ROOT / "data" / "synthetic" / "demo-game-dev-001_实验报告.docx")
        )

        def response_payload(evidence_id: str) -> str:
            return json.dumps(
                {
                    "criteria": [
                        {
                            "criterion_id": item["id"],
                            "score": item["max_score"] * 0.6,
                            "reason": "该维度有证据，但仍需教师终审。",
                            "evidence_ids": [evidence_id],
                            "confidence": 0.76,
                        }
                        for item in rubric["criteria"]
                    ]
                },
                ensure_ascii=False,
            )

        class RepairingClient:
            calls = 0
            requests = []

            def create_message(self, **kwargs):
                self.calls += 1
                self.requests.append(kwargs)
                return LLMResponse(
                    text_blocks=[
                        TextBlock(
                            response_payload(
                                "invented-evidence" if self.calls == 1 else "p-0001"
                            )
                        )
                    ],
                    input_tokens=120,
                    output_tokens=220,
                )

        client = RepairingClient()
        _, run, _ = build_model_trace(
            parsed,
            trace_id="trace-repair",
            submission_alias="submission-repair",
            rubric=rubric,
            client=client,
            config=LLMConfig(
                api_key="test-only",
                model="MiniMax-M3",
                enable_thinking=False,
                failover_endpoints=[],
            ),
        )
        repair_request = json.dumps(client.requests[1], ensure_ascii=False)
        self.assertEqual(run["attempts"], 2)
        self.assertIn("引用了未知证据", repair_request)
        self.assertIn("重新输出完整 JSON", repair_request)

    def test_public_synthetic_json_has_no_common_sensitive_identifiers(self):
        for path in (ROOT / "data" / "synthetic").glob("*.json"):
            self.assertEqual(find_sensitive_data(path.read_text(encoding="utf-8")), [])

    def test_cross_domain_examples_have_distinct_auditable_scores(self):
        examples = {
            "demo-allergen-001_实验报告.docx": 68,
            "demo-game-dev-001_实验报告.docx": 75,
        }
        for filename, expected_score in examples.items():
            path = ROOT / "data" / "synthetic" / filename
            self.assertTrue(path.exists())
            parsed = parse_report(str(path))
            trace = build_demo_trace(
                parsed,
                trace_id=f"trace-{path.stem}",
                submission_alias=path.stem,
            )
            self.assertEqual(trace.model_total_score, expected_score)
            self.assertTrue(trace.needs_human_review)
            self.assertGreaterEqual(len(trace.evidence), 6)
            self.assertTrue(
                all(
                    item.evidence_id.startswith(("p-", "t-", "i-"))
                    for item in trace.evidence
                )
            )
            self.assertEqual(
                len({item.evidence_id for item in trace.evidence}),
                len(trace.evidence),
            )
            text = "\n".join(
                str(item.get("text", "")) for item in parsed.get("paragraphs") or []
            )
            self.assertEqual(find_sensitive_data(text), [])

    def test_synthetic_docx_contains_review_fields_and_comments(self):
        input_path = ROOT / "data" / "synthetic" / "demo-student-001_实验报告.docx"
        graded_path = (
            ROOT / "data" / "synthetic" / "demo-student-001_实验报告_批改示例.docx"
        )
        self.assertTrue(input_path.exists())
        self.assertTrue(graded_path.exists())
        document = Document(str(graded_path))
        all_text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ]
        )
        self.assertIn("13+16+20+10+11+4=74", all_text)
        self.assertIn("正式成绩由教师复核确认", all_text)
        self.assertEqual(find_sensitive_data(all_text), [])
        with zipfile.ZipFile(graded_path) as package:
            self.assertIn("word/comments.xml", package.namelist())
            comments = package.read("word/comments.xml").decode("utf-8")
            self.assertEqual(comments.count("<w:comment "), 5)

    def test_public_demo_api_runs_full_review_loop(self):
        app = FastAPI()
        app.include_router(demo_api.router)
        sample_path = ROOT / "data" / "synthetic" / "demo-allergen-001_实验报告.docx"
        with tempfile.TemporaryDirectory(prefix="labtrace-test-") as temp_dir:
            with patch.object(demo_api, "RUNTIME_DIR", Path(temp_dir)):
                demo_api._tasks.clear()
                client = TestClient(app)
                bootstrap = client.get("/labtrace-api/bootstrap")
                self.assertEqual(bootstrap.status_code, 200)
                self.assertEqual(len(bootstrap.json()["samples"]), 2)
                self.assertEqual(bootstrap.json()["mode"]["id"], "deterministic_demo")
                self.assertEqual(
                    bootstrap.json()["assignment_template"]["input_url"],
                    "labtrace-api/sample/assignment-template",
                )
                for kind in ("assignment-template", "allergen", "game-dev"):
                    sample = client.get(f"/labtrace-api/sample/{kind}")
                    self.assertEqual(sample.status_code, 200)
                    self.assertTrue(sample.content.startswith(b"PK"))
                assignment = Document(
                    str(
                        ROOT
                        / "data"
                        / "synthetic"
                        / "demo-assignment-template_实验任务书.docx"
                    )
                )
                assignment_text = "\n".join(
                    paragraph.text for paragraph in assignment.paragraphs
                )
                self.assertIn("Unity 弹射原型", assignment_text)
                self.assertIn("人工合成公开模板", assignment_text)
                self.assertEqual(find_sensitive_data(assignment_text), [])
                rubric_template = client.get("/labtrace-api/sample/rubric")
                self.assertEqual(rubric_template.status_code, 200)
                self.assertEqual(
                    client.post(
                        "/labtrace-api/rubrics/validate",
                        json=json.loads(rubric_template.content),
                    ).status_code,
                    200,
                )
                with sample_path.open("rb") as stream:
                    response = client.post(
                        "/labtrace-api/grade",
                        files={
                            "report": (
                                sample_path.name,
                                stream,
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                        },
                    )
                self.assertEqual(response.status_code, 200, response.text)
                task = response.json()
                self.assertEqual(task["status"], "review_pending")
                self.assertTrue(task["trace"]["needs_human_review"])
                self.assertGreaterEqual(len(task["trace"]["evidence"]), 6)
                self.assertGreaterEqual(
                    task["delivery"]["details"]["annotations_count"], 4
                )
                self.assertTrue(
                    task["delivery"]["details"]["evidence_appendix_appended"]
                )
                self.assertEqual(
                    [item["reference_number"] for item in task["evidence_appendix"]],
                    list(range(1, len(task["evidence_appendix"]) + 1)),
                )
                self.assertTrue(
                    all(
                        item["location_label"].startswith("Word ")
                        for item in task["evidence_appendix"]
                    )
                )
                self.assertGreaterEqual(len(task["word_comments"]), 4)
                self.assertEqual(
                    len(task["word_comments"]),
                    task["word_workflow"]["native_comments"],
                )
                self.assertTrue(task["word_comments"][0]["reference_numbers"])
                criteria = []
                for item in task["trace"]["criteria"]:
                    if item["criterion_id"] == "analysis_and_validation":
                        item["score"] = 10
                    criteria.append(
                        {
                            "criterion_id": item["criterion_id"],
                            "score": item["score"],
                            "reason": item["reason"],
                        }
                    )
                reviewed = client.post(
                    "/labtrace-api/review",
                    json={
                        "task_id": task["task_id"],
                        "criteria": criteria,
                        "note": "教师已核对 ELISA 证据并确认两分调整。",
                    },
                )
                self.assertEqual(reviewed.status_code, 200, reviewed.text)
                completed = reviewed.json()
                self.assertEqual(completed["status"], "completed")
                self.assertEqual(completed["trace"]["review"]["status"], "adjusted")
                self.assertEqual(completed["trace"]["review"]["final_score"], 70)
                self.assertEqual(completed["diagnosis"]["record_count"], 4)
                self.assertEqual(completed["diagnosis"]["class_average"], 73.5)
                self.assertEqual(
                    completed["diagnosis"]["top_weaknesses"][0]["criterion_id"],
                    "analysis_and_validation",
                )
                report_download = client.get(
                    f"/labtrace-api/tasks/{task['task_id']}/download?kind=report"
                )
                self.assertEqual(report_download.status_code, 200)
                reviewed_document = Document(io.BytesIO(report_download.content))
                reviewed_text = "\n".join(
                    [paragraph.text for paragraph in reviewed_document.paragraphs]
                    + [
                        cell.text
                        for table in reviewed_document.tables
                        for row in table.rows
                        for cell in row.cells
                    ]
                )
                self.assertIn("教师已核对 ELISA 证据并确认两分调整", reviewed_text)
                self.assertIn("附录：证据引用索引（LabTrace）", reviewed_text)
                self.assertIn("[1]", reviewed_text)
                with zipfile.ZipFile(io.BytesIO(report_download.content)) as package:
                    self.assertIn("word/comments.xml", package.namelist())
                    comments = package.read("word/comments.xml").decode("utf-8")
                    self.assertGreaterEqual(comments.count("<w:comment "), 4)
                trace_download = client.get(
                    f"/labtrace-api/tasks/{task['task_id']}/download?kind=trace"
                )
                self.assertEqual(trace_download.status_code, 200)
                trace_payload = trace_download.json()
                self.assertEqual(
                    len(trace_payload["evidence_appendix"]),
                    len(completed["trace"]["evidence"]),
                )
                self.assertTrue(trace_payload["word_comments"])
                source_download = client.get(
                    f"/labtrace-api/tasks/{task['task_id']}/download?kind=source"
                )
                self.assertEqual(source_download.status_code, 200)
                self.assertEqual(source_download.content, sample_path.read_bytes())
                deleted = client.delete(f"/labtrace-api/tasks/{task['task_id']}")
                self.assertEqual(deleted.status_code, 200)
                self.assertEqual(
                    client.get(f"/labtrace-api/tasks/{task['task_id']}").status_code,
                    404,
                )

    def test_public_api_rejects_disguised_uploads_and_hides_local_paths(self):
        app = FastAPI()
        app.include_router(demo_api.router)
        sample_path = ROOT / "data" / "synthetic" / "demo-allergen-001_实验报告.docx"
        with tempfile.TemporaryDirectory(prefix="labtrace-security-") as temp_dir:
            with patch.object(demo_api, "RUNTIME_DIR", Path(temp_dir)):
                client = TestClient(app)
                invalid_pdf = client.post(
                    "/labtrace-api/grade",
                    files={"report": ("fake.pdf", b"not a pdf", "application/pdf")},
                )
                self.assertEqual(invalid_pdf.status_code, 400)
                invalid_docx = client.post(
                    "/labtrace-api/grade",
                    files={
                        "report": (
                            "fake.docx",
                            b"not a docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    },
                )
                self.assertEqual(invalid_docx.status_code, 400)
                with sample_path.open("rb") as stream:
                    response = client.post(
                        "/labtrace-api/grade",
                        files={
                            "report": (
                                sample_path.name,
                                stream,
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                        },
                    )
                self.assertEqual(response.status_code, 200, response.text)
                payload = response.json()
                self.assertRegex(payload["task_id"], r"^[0-9a-f]{32}$")
                serialized = json.dumps(payload, ensure_ascii=False)
                self.assertNotIn("input_path", serialized)
                self.assertNotIn("output_path", serialized)
                self.assertNotIn(temp_dir, serialized)


if __name__ == "__main__":
    unittest.main()
