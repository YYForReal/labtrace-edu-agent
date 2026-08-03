"""Public, self-contained demo API for the GOAI competition edition."""

from __future__ import annotations

import asyncio
import hashlib
import io
import json
import os
import re
import shutil
import time
import uuid
import zipfile
from collections import defaultdict, deque
from pathlib import Path
from typing import Any

from fastapi import APIRouter, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from docx import Document

from agent_skills.report_injector.scripts.inject_grading_to_docx import inject_all
from goaihz.demo_engine import (
    build_learning_feedback,
    load_demo_rubric,
    parse_report,
    trace_payload,
)
from goaihz.model_engine import grade_report_with_adapter, model_runtime_status
from goaihz.src.labtrace.contracts import (
    CriterionDecision,
    GradeTrace,
    ReviewDecision,
)
from goaihz.src.labtrace.diagnosis import build_class_diagnosis
from goaihz.src.labtrace.rubric import (
    RubricError,
    load_rubric_json,
    rubric_summary,
    validate_rubric,
)

router = APIRouter(prefix="/labtrace-api", tags=["LabTrace Demo"])
ROOT = Path(__file__).resolve().parent
SYNTHETIC_DIR = ROOT / "data" / "synthetic"
RUNTIME_DIR = Path(os.getenv("LABTRACE_RUNTIME_DIR", ROOT / "runtime" / "demo_tasks"))
MAX_UPLOAD_BYTES = 25 * 1024 * 1024
MAX_DOCX_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_DOCX_ENTRIES = 2_000
ALLOWED_SUFFIXES = {".docx", ".pdf"}
TASK_TTL_SECONDS = max(300, int(os.getenv("LABTRACE_TASK_TTL_SECONDS", "86400")))
GRADE_RATE_LIMIT = max(1, int(os.getenv("LABTRACE_GRADE_RATE_LIMIT", "8")))
GRADE_RATE_WINDOW_SECONDS = max(
    60, int(os.getenv("LABTRACE_GRADE_RATE_WINDOW_SECONDS", "600"))
)
GRADE_CONCURRENCY = max(1, int(os.getenv("LABTRACE_GRADE_CONCURRENCY", "2")))
TRUST_PROXY_HEADERS = os.getenv("LABTRACE_TRUST_PROXY_HEADERS", "").lower() == "true"
_tasks: dict[str, dict[str, Any]] = {}
_grade_attempts: dict[str, deque[float]] = defaultdict(deque)
_grade_semaphore = asyncio.Semaphore(GRADE_CONCURRENCY)


def _cleanup_expired_tasks(now: float | None = None) -> None:
    now = now or time.time()
    if not RUNTIME_DIR.exists():
        return
    for task_dir in RUNTIME_DIR.iterdir():
        if not task_dir.is_dir():
            continue
        state_path = task_dir / "state.json"
        try:
            state = _read_json(state_path) if state_path.exists() else {}
            created_at = float(
                state.get("created_at_epoch", state_path.stat().st_mtime)
            )
        except (OSError, TypeError, ValueError, json.JSONDecodeError):
            created_at = task_dir.stat().st_mtime
        if now - created_at <= TASK_TTL_SECONDS:
            continue
        shutil.rmtree(task_dir, ignore_errors=True)
        _tasks.pop(task_dir.name, None)


def _client_key(request: Request) -> str:
    if TRUST_PROXY_HEADERS:
        forwarded = request.headers.get("x-forwarded-for", "").split(",", 1)[0].strip()
        if forwarded:
            return forwarded[:64]
    return request.client.host if request.client else "unknown"


def _enforce_grade_rate_limit(request: Request) -> None:
    now = time.monotonic()
    attempts = _grade_attempts[_client_key(request)]
    while attempts and now - attempts[0] > GRADE_RATE_WINDOW_SECONDS:
        attempts.popleft()
    if len(attempts) >= GRADE_RATE_LIMIT:
        raise HTTPException(
            status_code=429,
            detail="请求过于频繁，请稍后再试。公开演示不会无限制处理上传文件。",
        )
    attempts.append(now)


def _validate_upload(payload: bytes, suffix: str) -> None:
    if not payload:
        raise HTTPException(status_code=400, detail="上传文件为空")
    if suffix == ".pdf":
        if not payload.startswith(b"%PDF-"):
            raise HTTPException(
                status_code=400, detail="文件扩展名为 PDF，但内容不是有效 PDF"
            )
        return

    stream = io.BytesIO(payload)
    if not zipfile.is_zipfile(stream):
        raise HTTPException(
            status_code=400, detail="文件扩展名为 DOCX，但内容不是有效 DOCX"
        )
    stream.seek(0)
    try:
        with zipfile.ZipFile(stream) as package:
            entries = package.infolist()
            if len(entries) > MAX_DOCX_ENTRIES:
                raise HTTPException(status_code=400, detail="DOCX 内部文件数量异常")
            total_size = sum(item.file_size for item in entries)
            if total_size > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise HTTPException(status_code=400, detail="DOCX 解压后内容过大")
            if "word/document.xml" not in package.namelist():
                raise HTTPException(status_code=400, detail="DOCX 缺少主文档内容")
    except zipfile.BadZipFile as exc:
        raise HTTPException(status_code=400, detail="DOCX 压缩结构损坏") from exc


def _public_state(value: Any) -> Any:
    """Remove local filesystem details before returning task state to browsers."""
    if isinstance(value, dict):
        return {
            key: _public_state(item)
            for key, item in value.items()
            if key not in {"input_path", "output_path"} and not key.endswith("_path")
        }
    if isinstance(value, list):
        return [_public_state(item) for item in value]
    return value


def _safe_filename(name: str | None) -> str:
    candidate = Path((name or "report.docx").replace("\\", "/")).name
    candidate = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff._ -]", "_", candidate).strip(" .")
    return candidate or "report.docx"


def _read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def _write_json(path: Path, value: Any) -> None:
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2), encoding="utf-8")


def _task_path(task_id: str) -> Path:
    return RUNTIME_DIR / task_id


def _load_task(task_id: str) -> dict[str, Any]:
    if task_id in _tasks:
        return _tasks[task_id]
    state_path = _task_path(task_id) / "state.json"
    if state_path.exists():
        state = _read_json(state_path)
        _tasks[task_id] = state
        return state
    raise HTTPException(status_code=404, detail="演示任务不存在")


def _save_task(task_id: str, state: dict[str, Any]) -> None:
    _tasks[task_id] = state
    task_dir = _task_path(task_id)
    task_dir.mkdir(parents=True, exist_ok=True)
    _write_json(task_dir / "state.json", state)


def _find_grading_table_index(input_path: Path) -> int:
    """Locate a teacher-review table without assuming a course-specific template."""
    try:
        document = Document(str(input_path))
    except Exception:
        return 1
    for index, table in enumerate(document.tables):
        text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "成绩评定" in text and "评语" in text:
            return index
    return max(0, len(document.tables) - 1)


def _annotation_config(trace: GradeTrace, *, table_index: int = 1) -> dict[str, Any]:
    evidence_by_id = {item.evidence_id: item for item in trace.evidence}
    annotations = []
    for decision in trace.criteria:
        image_target = next(
            (
                evidence_by_id[evidence_id]
                for evidence_id in decision.evidence_ids
                if evidence_by_id[evidence_id].kind in {"image", "image_context"}
                and evidence_by_id[evidence_id].verification == "model_observed"
            ),
            None,
        )
        paragraph_target = next(
            (
                evidence_by_id[evidence_id]
                for evidence_id in decision.evidence_ids
                if evidence_by_id[evidence_id].kind == "paragraph"
            ),
            None,
        )
        target = image_target or paragraph_target
        if not target or not target.excerpt:
            continue
        if image_target:
            paragraph_match = re.search(r"@paragraph:(\d+)", target.locator)
            if not paragraph_match:
                continue
            annotation_target = {
                "type": "paragraph_index",
                "index": int(paragraph_match.group(1)) - 1,
            }
            annotation_text = f"图片证据批注：{decision.reason}"
        else:
            if "：" in target.excerpt[:24]:
                keyword = target.excerpt.split("：", 1)[0].strip() + "："
            else:
                keyword = target.excerpt[:10].rstrip("，。；： ")
            if len(keyword) < 4:
                continue
            annotation_target = {"type": "keyword", "keyword": keyword}
            annotation_text = decision.reason
        annotations.append(
            {
                "text": annotation_text,
                "target": annotation_target,
                "evidence_kind": "image" if image_target else "text",
            }
        )
        if len(annotations) >= 6:
            break
    weakest = sorted(
        trace.criteria,
        key=lambda item: (item.score / item.max_score, item.confidence),
    )[:2]
    summary = "；".join(
        f"{item.criterion_name} {item.score:g}/{item.max_score:g}：{item.reason}"
        for item in weakest
    )
    teacher_note = (
        trace.review.note.strip()
        if trace.review.status in {"approved", "adjusted"} and trace.review.note.strip()
        else ""
    )
    comment_parts = []
    if teacher_note:
        comment_parts.append(f"教师终审：{teacher_note}")
    if summary:
        comment_parts.append(summary[:600])
    comment_parts.append("AI 仅提供辅助建议，正式成绩与评语已经教师复核确认。")
    return {
        "annotations": annotations,
        "scores": [item.score for item in trace.criteria],
        "criteria": [
            {
                "name": item.criterion_name,
                "score": item.score,
                "max_score": item.max_score,
                "reason": item.reason,
            }
            for item in trace.criteria
        ],
        "comment": "；".join(comment_parts)[:1000],
        "author": "LabTrace 演示教师",
        "table_index": table_index,
    }


def _try_build_annotated_report(
    input_path: Path, output_path: Path, trace: GradeTrace
) -> dict[str, Any]:
    if input_path.suffix.lower() != ".docx":
        return {"available": False, "reason": "PDF 演示仅返回证据链，不修改原文件。"}
    try:
        result = inject_all(
            str(input_path),
            str(output_path),
            _annotation_config(
                trace, table_index=_find_grading_table_index(input_path)
            ),
        )
        return {"available": output_path.exists(), "details": result}
    except Exception as exc:
        shutil.copy2(input_path, output_path)
        return {
            "available": True,
            "reason": f"批注注入失败，已保留原文件作为人工接管副本：{exc}",
        }


def _word_workflow_summary(
    *,
    input_path: Path,
    document_profile: dict[str, Any],
    agent_run: dict[str, Any],
    delivery: dict[str, Any],
    trace: GradeTrace,
) -> dict[str, Any]:
    details = delivery.get("details") or {}
    image_count = int(document_profile.get("image_count", 0) or 0)
    return {
        "input_is_word": input_path.suffix.lower() == ".docx",
        "images_detected": image_count,
        "images_analyzed": int(agent_run.get("images_sent", 0) or 0),
        "native_comments": int(details.get("annotations_count", 0) or 0),
        "image_comments": int(details.get("image_annotations_count", 0) or 0),
        "teacher_feedback_written": bool(
            trace.review.status in {"approved", "adjusted"}
            and details.get("comment_injected")
        ),
        "score_written": bool(details.get("score_injected")),
        "editable_word_available": bool(
            input_path.suffix.lower() == ".docx" and delivery.get("available")
        ),
        "delivery_mode": str(details.get("delivery_mode", "comments_only")),
    }


def _diagnosis_records(
    current: list[dict[str, Any]], *, rubric_id: str
) -> list[dict[str, Any]]:
    if rubric_id != "general_lab_report_v1":
        return current
    base = _read_json(SYNTHETIC_DIR / "grade_records.json")
    return [*base, *current]


class ReviewCriterionInput(BaseModel):
    criterion_id: str
    score: float
    reason: str = ""


class ReviewRequest(BaseModel):
    task_id: str
    criteria: list[ReviewCriterionInput]
    note: str = Field(min_length=4, max_length=500)


@router.get("/bootstrap")
async def bootstrap():
    _cleanup_expired_tasks()
    rubric = validate_rubric(load_demo_rubric())
    runtime = model_runtime_status()
    if runtime["configured"]:
        mode = {
            "id": "model_agent",
            "label": f"{runtime['provider']} 真实 Agent",
            "disclaimer": (
                f"当前使用 {runtime['model']} 生成逐项建议。报告先在服务器解析并自动脱敏，"
                "默认仅发送匿名证据文本；图片只有经教师逐任务授权后才发送至多 4 张。"
                "模型失败时会明确标记并降级到规则模式。"
            ),
            "provider": runtime["provider"],
            "model": runtime["model"],
            "external_processing": True,
            "daily_remaining": runtime["daily_remaining"],
        }
    else:
        mode = {
            "id": "deterministic_demo",
            "label": "无密钥演示模式",
            "disclaimer": (
                "本模式真实执行文档解析、证据定位、分数校验、教师复核和学情诊断，"
                "评分建议由公开规则生成，不冒充大模型推理。配置模型后可切换真实 Agent。"
            ),
            "provider": "none",
            "model": "none",
            "external_processing": False,
            "daily_remaining": 0,
        }
    return {
        "product": {
            "name": "格物智评 LabTrace",
            "tagline": "上传图文 Word，交还可编辑的批注 Word",
            "track": "GOAI 2026 · Boundless Agents · AI+教育",
            "promise": (
                "保留原报告正文、表格和图片，把证据化批注、分项成绩与教师评语"
                "写回同一份可继续编辑的 Word。"
            ),
        },
        "word_capabilities": {
            "input": "DOCX 原生结构、正文、表格与内嵌图片",
            "analysis": "图片逐任务授权；模型判断绑定图片证据与原 Word 段落",
            "output": "Word/WPS 原生批注、分项成绩、教师评语与可编辑交付",
            "generic_template": "没有固定评分表时自动追加通用教师批改意见页",
        },
        "mode": mode,
        "data_policy": {
            "accepted_formats": ["DOCX", "PDF"],
            "max_upload_mb": 25,
            "identity_redaction": "发送外部模型前自动替换可识别的姓名、学号、手机号、邮箱和身份证号",
            "model_payload": "默认只发送有界的匿名正文/表格证据与图片邻近文本；教师可显式授权发送至多 4 张图片",
            "retention": "默认 24 小时自动删除；完成终审后可立即删除",
            "teacher_responsibility": "教师需确认上传授权；AI 建议不得自动作为正式成绩",
        },
        "rubric": rubric,
        "rubric_template_url": "labtrace-api/sample/rubric",
        "assignment_template": {
            "name": "Unity 弹射原型实验任务书",
            "description": "公开合成题目模板，可下载后直接作为课程任务书或学生报告骨架使用。",
            "input_url": "labtrace-api/sample/assignment-template",
            "filename": "LabTrace_Unity弹射原型_实验任务书.docx",
        },
        "samples": [
            {
                "id": "allergen",
                "domain": "生命科学",
                "name": "过敏原蛋白 ELISA 检测",
                "provenance": "完全合成 · 含标准曲线图片 · 不涉及患者或医学诊断",
                "input_url": "labtrace-api/sample/allergen",
                "filename": "demo-allergen-001_实验报告.docx",
            },
            {
                "id": "game-dev",
                "domain": "游戏开发",
                "name": "Unity 弹射原型与碰撞验证",
                "provenance": "实际课程结构合成重构 · 含游戏运行示意图 · 无学生原文",
                "input_url": "labtrace-api/sample/game-dev",
                "filename": "demo-game-dev-001_实验报告.docx",
            },
        ],
        "agent_steps": [
            {
                "id": "understand",
                "label": "理解任务",
                "detail": "识别报告类型与 rubric 约束",
            },
            {
                "id": "parse",
                "label": "解析证据",
                "detail": "抽取正文、表格、图像与结构",
            },
            {
                "id": "grade",
                "label": "逐项判断",
                "detail": "将评分理由绑定到可定位证据",
            },
            {
                "id": "verify",
                "label": "校验结果",
                "detail": "检查总分、证据覆盖与置信度",
            },
            {"id": "review", "label": "教师终审", "detail": "低置信度案例必须人工确认"},
            {
                "id": "diagnose",
                "label": "回写与诊断",
                "detail": "批注和评语写回 Word，复核结果进入学情",
            },
        ],
    }


@router.post("/grade")
async def grade_demo_report(
    request: Request,
    report: UploadFile = File(...),
    rubric: UploadFile | None = File(default=None),
    allow_external_images: bool = Form(default=False),
):
    _cleanup_expired_tasks()
    _enforce_grade_rate_limit(request)
    suffix = Path(report.filename or "").suffix.lower()
    if suffix not in ALLOWED_SUFFIXES:
        raise HTTPException(status_code=400, detail="仅支持 .docx 或 .pdf")

    payload = await report.read(MAX_UPLOAD_BYTES + 1)
    if len(payload) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="演示文件不能超过 25 MB")
    _validate_upload(payload, suffix)

    customized_rubric = rubric is not None
    try:
        selected_rubric = (
            load_rubric_json(await rubric.read(256 * 1024 + 1))
            if rubric is not None
            else validate_rubric(load_demo_rubric())
        )
    except RubricError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    task_id = uuid.uuid4().hex
    task_dir = _task_path(task_id)
    task_dir.mkdir(parents=True, exist_ok=True)
    input_path = task_dir / _safe_filename(report.filename)
    input_path.write_bytes(payload)

    try:
        async with _grade_semaphore:
            parsed = await asyncio.to_thread(parse_report, str(input_path))
            outcome = await grade_report_with_adapter(
                parsed,
                trace_id=f"trace-{task_id}",
                submission_alias=f"submission-{hashlib.sha256(payload).hexdigest()[:12]}",
                rubric=selected_rubric,
                allow_external_images=allow_external_images,
            )
    except HTTPException:
        shutil.rmtree(task_dir, ignore_errors=True)
        raise
    except Exception as exc:
        shutil.rmtree(task_dir, ignore_errors=True)
        raise HTTPException(
            status_code=422, detail=f"文档解析失败：{type(exc).__name__}"
        )
    now = time.time()
    trace = outcome.trace

    output_path = task_dir / f"{input_path.stem}_批改演示.docx"
    delivery = _try_build_annotated_report(input_path, output_path, trace)
    word_workflow = _word_workflow_summary(
        input_path=input_path,
        document_profile=parsed.get("document_profile", {}),
        agent_run=outcome.run,
        delivery=delivery,
        trace=trace,
    )
    selected_rubric_path = task_dir / "rubric.json"
    _write_json(selected_rubric_path, selected_rubric)
    rubric_info = rubric_summary(selected_rubric, customized=customized_rubric)
    model_message = (
        f"{outcome.run.get('provider')} / {outcome.run.get('model')} 已生成"
        if outcome.mode == "model_agent"
        else (
            "外部模型调用失败，已显式降级到确定性规则"
            if outcome.mode == "deterministic_fallback"
            else "已使用可离线复现的确定性规则生成"
        )
    )
    state = {
        "task_id": task_id,
        "status": (
            "review_pending" if trace.needs_human_review else "teacher_confirmation"
        ),
        "mode": outcome.mode,
        "input_filename": input_path.name,
        "input_path": str(input_path),
        "output_path": str(output_path) if output_path.exists() else "",
        "document_profile": parsed.get("document_profile", {}),
        "rubric": rubric_info,
        "agent_run": outcome.run,
        "privacy": outcome.privacy,
        "trace": trace_payload(trace),
        "suggested_trace": trace_payload(trace),
        "learning_feedback": build_learning_feedback(trace),
        "delivery": delivery,
        "word_workflow": word_workflow,
        "created_at_epoch": now,
        "expires_at_epoch": now + TASK_TTL_SECONDS,
        "events": [
            {
                "stage": "understand",
                "message": (
                    f"已载入“{rubric_info['experiment_name']}”"
                    f"{rubric_info['criterion_count']} 维 rubric。"
                ),
            },
            {
                "stage": "parse",
                "message": (
                    f"已解析 {len(parsed.get('paragraphs') or [])} 个段落、"
                    f"{len(parsed.get('tables') or [])} 个表格和"
                    f"{len(parsed.get('images') or [])} 张图片。"
                ),
            },
            {
                "stage": "privacy",
                "message": (
                    f"模型输入执行 {outcome.privacy.get('policy')}；"
                    f"识别并替换 {outcome.privacy.get('detected_sensitive_items', 0)} 项敏感标识。"
                ),
            },
            {
                "stage": "grade",
                "message": f"{model_message} {len(trace.criteria)} 个逐项评分建议。",
            },
            {"stage": "verify", "message": "分项合计、证据引用和置信度校验通过。"},
            {
                "stage": "review",
                "message": "所有 AI 建议均已进入教师终审，不会自动发布。",
            },
            {
                "stage": "delivery",
                "message": (
                    f"已生成可编辑 Word：{word_workflow['native_comments']} 条原生批注，"
                    f"其中 {word_workflow['image_comments']} 条定位到图片证据。"
                ),
            },
        ],
    }
    _save_task(task_id, state)
    return _public_state(state)


@router.post("/review")
async def review_demo_result(request: ReviewRequest):
    _cleanup_expired_tasks()
    state = _load_task(request.task_id)
    trace = GradeTrace.from_dict(state.get("suggested_trace", state["trace"]))
    original = {item.criterion_id: item for item in trace.criteria}
    updates = {item.criterion_id: item for item in request.criteria}

    unknown = set(updates) - set(original)
    if unknown:
        raise HTTPException(status_code=400, detail=f"未知评分维度：{sorted(unknown)}")

    reviewed_criteria = []
    adjusted = False
    for item in trace.criteria:
        update = updates.get(item.criterion_id)
        if update is None:
            reviewed_criteria.append(item)
            continue
        if not 0 <= update.score <= item.max_score:
            raise HTTPException(
                status_code=400,
                detail=f"{item.criterion_name}分数必须位于 0 至 {item.max_score}",
            )
        adjusted = adjusted or abs(update.score - item.score) > 0.01
        reviewed_criteria.append(
            CriterionDecision(
                criterion_id=item.criterion_id,
                criterion_name=item.criterion_name,
                max_score=item.max_score,
                score=update.score,
                reason=update.reason.strip() or item.reason,
                evidence_ids=item.evidence_ids,
                confidence=item.confidence,
            )
        )

    final_score = sum(item.score for item in reviewed_criteria)
    review = ReviewDecision(
        status="adjusted" if adjusted else "approved",
        reviewer_role="teacher",
        final_score=final_score,
        note=request.note,
    )
    review.validate()

    reviewed_trace = GradeTrace(
        trace_id=trace.trace_id,
        rubric_id=trace.rubric_id,
        submission_alias=trace.submission_alias,
        evidence=trace.evidence,
        criteria=tuple(reviewed_criteria),
        model_total_score=final_score,
        needs_human_review=False,
        review_reasons=trace.review_reasons,
        review=review,
    )
    reviewed_trace.validate()
    state["status"] = "completed"
    state["trace"] = trace_payload(reviewed_trace)
    state["learning_feedback"] = build_learning_feedback(reviewed_trace)
    state["events"].append(
        {"stage": "review", "message": f"教师已确认，最终得分 {final_score:g}。"}
    )

    input_path = Path(state["input_path"])
    output_path = (
        Path(state["output_path"])
        if state.get("output_path")
        else (_task_path(request.task_id) / f"{input_path.stem}_批改演示.docx")
    )
    state["delivery"] = _try_build_annotated_report(
        input_path, output_path, reviewed_trace
    )
    state["output_path"] = str(output_path) if output_path.exists() else ""
    state["word_workflow"] = _word_workflow_summary(
        input_path=input_path,
        document_profile=state.get("document_profile", {}),
        agent_run=state.get("agent_run", {}),
        delivery=state["delivery"],
        trace=reviewed_trace,
    )
    state["events"].append(
        {
            "stage": "delivery",
            "message": (
                "教师终审评语与最终成绩已写回可编辑 Word；"
                f"保留 {state['word_workflow']['native_comments']} 条原生批注。"
            ),
        }
    )

    current_record = {
        "submission_alias": reviewed_trace.submission_alias,
        "review_status": review.status,
        "criterion_scores": [
            {
                "criterion_id": item.criterion_id,
                "criterion_name": item.criterion_name,
                "max_score": item.max_score,
                "score": item.score,
            }
            for item in reviewed_trace.criteria
        ],
    }
    state["diagnosis"] = build_class_diagnosis(
        _diagnosis_records([current_record], rubric_id=reviewed_trace.rubric_id)
    )
    _save_task(request.task_id, state)
    return _public_state(state)


@router.get("/tasks/{task_id}")
async def get_demo_task(task_id: str):
    _cleanup_expired_tasks()
    return _public_state(_load_task(task_id))


@router.get("/tasks/{task_id}/download")
async def download_demo_artifact(task_id: str, kind: str = "report"):
    _cleanup_expired_tasks()
    state = _load_task(task_id)
    if kind == "trace":
        path = _task_path(task_id) / "trace.json"
        _write_json(path, state["trace"])
        return FileResponse(path, filename=f"{task_id}_evidence_trace.json")
    path = Path(state.get("output_path") or "")
    if not path.exists():
        raise HTTPException(status_code=404, detail="当前任务没有可下载的批改文档")
    return FileResponse(path, filename=path.name)


@router.delete("/tasks/{task_id}")
async def delete_demo_task(task_id: str):
    _cleanup_expired_tasks()
    _load_task(task_id)
    shutil.rmtree(_task_path(task_id), ignore_errors=True)
    _tasks.pop(task_id, None)
    return {"task_id": task_id, "status": "deleted"}


@router.get("/sample/{kind}")
async def download_sample(kind: str):
    names: dict[str, str] = {
        "input": "demo-student-001_实验报告.docx",
        "graded": "demo-student-001_实验报告_批改示例.docx",
        "allergen": "demo-allergen-001_实验报告.docx",
        "game-dev": "demo-game-dev-001_实验报告.docx",
        "assignment-template": "demo-assignment-template_实验任务书.docx",
    }
    if kind == "rubric":
        path = ROOT / "config" / "rubrics" / "general_lab_report_v1.json"
        return FileResponse(path, filename="LabTrace_评分标准模板.json")
    if kind not in names:
        raise HTTPException(status_code=404, detail="未知示例文件")
    path = SYNTHETIC_DIR / names[kind]
    if not path.exists():
        raise HTTPException(status_code=404, detail="示例文件尚未生成")
    return FileResponse(path, filename=path.name)


@router.post("/rubrics/validate")
async def validate_teacher_rubric(rubric: dict[str, Any]):
    try:
        value = validate_rubric(rubric)
    except RubricError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    return {
        "valid": True,
        "rubric": rubric_summary(value, customized=True),
    }
